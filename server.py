from __future__ import annotations

import json
import math
import mimetypes
import os
import re
import warnings
from dataclasses import asdict, dataclass, field
from datetime import UTC, date, datetime, timedelta
from http.server import BaseHTTPRequestHandler, ThreadingHTTPServer
from io import BytesIO
from pathlib import Path
from typing import Any
from urllib.parse import parse_qs, urlparse

with warnings.catch_warnings():
    warnings.simplefilter("ignore", DeprecationWarning)
    import cgi

from openpyxl import load_workbook

try:
    from docx import Document
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.shared import Pt
except ImportError:  # pragma: no cover - optional dependency for docx export
    Document = None
    WD_ALIGN_PARAGRAPH = None
    Pt = None

try:
    from reportlab.lib.pagesizes import A4
    from reportlab.lib.units import mm
    from reportlab.pdfgen import canvas as pdf_canvas
except ImportError:  # pragma: no cover - optional dependency for pdf export
    A4 = None
    mm = None
    pdf_canvas = None


ROOT = Path(__file__).resolve().parent
STATIC_DIR = ROOT / "static"
QUALITY_REPORTS_DIR = ROOT / "uploads" / "quality_reports"
QUALITY_REPORTS_DIR.mkdir(parents=True, exist_ok=True)

KNOWLEDGE_BASE_DIR = ROOT / "uploads" / "knowledge_base"
KNOWLEDGE_BASE_DIR.mkdir(parents=True, exist_ok=True)

CASE_ATTACHMENTS_DIR = ROOT / "uploads" / "case_attachments"
CASE_ATTACHMENTS_DIR.mkdir(parents=True, exist_ok=True)

SHIPPER_NOMINATIONS_DIR = ROOT / "uploads" / "shipper_nominations"
SHIPPER_NOMINATIONS_DIR.mkdir(parents=True, exist_ok=True)


ROLES = ["TRANSPORTER", "SHIPPER", "SUPPLIER", "CUSTOMER", "GAS_MARKETER", "REGULATOR"]

# Knowledge Base — categories for uploaded policy, regulatory, and reference documents.
KNOWLEDGE_CATEGORIES = ["REGULATION", "POLICY", "SOP", "CIRCULAR", "TEMPLATE", "OTHER"]
KNOWLEDGE_CATEGORY_LABELS = {
    "REGULATION": "Act / Regulation",
    "POLICY": "Policy",
    "SOP": "SOP / Guideline",
    "CIRCULAR": "Circular / Directive",
    "TEMPLATE": "Template / Form",
    "OTHER": "Other Reference",
}

# Case Management — complaint categories, complainant roles, and priorities for
# manually logged complaints (as distinct from AI-flagged exceptions, which route
# through the Escalations workflow instead).
COMPLAINT_CATEGORIES = [
    "BILLING_ATTRIBUTION",
    "GAS_QUALITY",
    "METERING_MLF",
    "PRESSURE_DELIVERY",
    "CONTRACTUAL_NEA",
    "SERVICE_COMPLAINT",
    "OTHER",
]
COMPLAINT_CATEGORY_LABELS = {
    "BILLING_ATTRIBUTION": "Billing / Attribution Dispute",
    "GAS_QUALITY": "Gas Quality",
    "METERING_MLF": "Metering / MLF Dispute",
    "PRESSURE_DELIVERY": "Pressure / Delivery Shortfall",
    "CONTRACTUAL_NEA": "Contractual / NEA",
    "SERVICE_COMPLAINT": "Service Complaint",
    "OTHER": "Other",
}
COMPLAINANT_ROLES = ["SHIPPER", "SUPPLIER", "TRANSPORTER", "CUSTOMER", "PUBLIC"]
COMPLAINT_PRIORITIES = ["LOW", "MEDIUM", "HIGH", "CRITICAL"]
COMPLAINT_STATUSES = ["OPEN", "IN_PROGRESS", "RESOLVED", "CLOSED"]

# Configurable monitoring bands. Percentages, e.g. 3.0 = +/-3%.
# NOTE: these are the defaults referenced throughout the dashboard and seed data.
# The demo data's built-in shrinkage/meter-linearity-factor/imbalance losses create a
# "normal" injection-vs-uptake gap of roughly 3-6%, so a literal +/-3% band (the figure
# often used as an illustrative example) would flag almost every day. These defaults are
# set a little above that normal noise band so flags stay exceptional; adjust to match
# whatever band NGIC/the regulator actually agrees on for a live network.
NOMINATION_VARIANCE_THRESHOLD_PCT = 5.0
LINE_PACK_VARIANCE_THRESHOLD_PCT = 5.5

# MLF (Meter Linearity / Measurement Loss Factor) is a multiplier close to 1.0 applied to
# raw injection to get effective injection (effective = injection * mlf). The demo data's
# MLF factors sit in the 0.9975-0.9985 range, i.e. a "normal" measurement loss of roughly
# 0.15-0.25%. This threshold flags days where that loss is unusually large, which can
# indicate a meter calibration issue rather than routine measurement variance.
MLF_LOSS_THRESHOLD_PCT = 0.4

# Escalation workflow stages, in order. Mirrors the administrative escalation path agreed
# with the regulator: an AI-flagged exception starts with the analyst and, if unresolved,
# moves up the chain toward possible regulatory/administrative action.
ESCALATION_STAGES = ["ANALYST", "MANAGER", "DIRECTOR", "DG", "ACE"]
ESCALATION_STAGE_LABELS = {
    "ANALYST": "Analyst",
    "MANAGER": "Manager",
    "DIRECTOR": "Director",
    "DG": "Director-General",
    "ACE": "Ag. Chief Executive",
}

# Recommended action text keyed by exception rule. Referenced by both the Alerts &
# Exceptions view and the auto-generated escalation report. Framed against the Nigerian
# Gas Transportation Network Code (NGTNC) balancing/administrative-action provisions, as
# discussed with the regulator; wording should be reviewed against the actual code text
# before this is used outside a demo context.
RECOMMENDED_ACTIONS = {
    "PRESSURE_OUTSIDE_NEA": "Request an operational explanation from the supplier for the entry pressure excursion; if recurring, recommend a maintenance/inspection review of the entry point metering skid.",
    "CONDENSATE_THRESHOLD": "Flag to the supplier for condensate handling review at the entry point; monitor for repeat breaches before escalating.",
    "SHRINKAGE_THRESHOLD": "Request supporting metering data from the transporter for the affected exit point; if shrinkage persists above threshold, recommend a joint meter proving exercise.",
    "NOMINATION_VARIANCE": "Issue an administrative query to the accountable shipper under the NGTNC nomination/balancing provisions requesting an explanation for the under- or over-delivery; cross-check linked offtaker withdrawals before further action.",
    "LINE_PACK_VARIANCE": "Refer to NGIC for confirmation of network line-pack status; if the gap is attributable to a specific shipper drawing beyond their supplier's delivered volume, recommend administrative action against that shipper under the NGTNC code.",
    "MLF_LOSS_THRESHOLD": "Recommend a meter calibration / linearity check at the entry point; measurement loss materially above the normal band understates deliverable volume and should be verified before it affects settlement.",
    "ATTRIBUTION_EXCEEDS_EFFECTIVE_INJECTION": "Recommend a roll-up reconciliation with the shipper, linked supplier/GASCO, and affected offtakers; attribution exceeding effective injection points to either a data entry error or unaccounted-for gas.",
}
DEFAULT_RECOMMENDED_ACTION = "Review the underlying submission with the supplier and confirm whether the variance is a data issue or an operational one before deciding on next steps."

# Regulatory citations referenced by the AI Consequence Management engine when it flags a
# potential incident. As-is, an analyst would manually search the repository of guidelines,
# SOPs, and codes for the relevant clause; the AI attaches the citation automatically so the
# escalation report/letter can quote a specific Act, section, and code reference. Wording is
# indicative pending sign-off from Legal/Regulatory Affairs against the current PIA/NGTNC text.
PIA_REFERENCES: dict[str, dict[str, str]] = {
    "PRESSURE_OUTSIDE_NEA": {"act": "Petroleum Industry Act 2021 (PIA)", "section": "Section 167 (Third Party Access to Facilities) / NEA operating conditions", "code": "NGTNC Part D - Entry Point Operating Conditions"},
    "CONDENSATE_THRESHOLD": {"act": "Petroleum Industry Act 2021 (PIA)", "section": "Section 167(3)(b) - technical/operational conditions of access", "code": "NGTNC Part D - Gas Quality and Condensate Handling"},
    "SHRINKAGE_THRESHOLD": {"act": "Petroleum Industry Act 2021 (PIA)", "section": "Section 166 (Regulations for gas transportation)", "code": "NGTNC Part E - Measurement, Shrinkage and Losses"},
    "NOMINATION_VARIANCE": {"act": "Petroleum Industry Act 2021 (PIA)", "section": "Section 166 - gas transportation network code compliance", "code": "NGTNC Part C - Nomination and Balancing"},
    "LINE_PACK_VARIANCE": {"act": "Petroleum Industry Act 2021 (PIA)", "section": "Section 166 - gas transportation network code compliance", "code": "NGTNC Part C - Balancing and Line-Pack Management"},
    "MLF_LOSS_THRESHOLD": {"act": "Petroleum Industry Act 2021 (PIA)", "section": "Section 166 - gas transportation network code compliance", "code": "NGTNC Part E - Measurement and Metering Standards"},
    "ATTRIBUTION_EXCEEDS_EFFECTIVE_INJECTION": {"act": "Petroleum Industry Act 2021 (PIA)", "section": "Section 166 - gas transportation network code compliance", "code": "NGTNC Part C - Allocation and Attribution"},
    "OFF_SPEC_GAS": {"act": "Petroleum Industry Act 2021 (PIA)", "section": "Section 167(3)(b) - technical/operational conditions of access", "code": "NGTNC Part D - Gas Quality Specification"},
}
DEFAULT_PIA_REFERENCE = {"act": "Petroleum Industry Act 2021 (PIA)", "section": "Section 166 - gas transportation network code compliance", "code": "Nigerian Gas Transportation Network Code (NGTNC)"}

RECOMMENDED_ACTIONS["OFF_SPEC_GAS"] = (
    "Reject/quarantine the affected volume pending confirmation, notify the supplier/GASCO of the off-spec "
    "lab result, and request a corrective quality report before the next injection window; escalate if the "
    "supplier delivers off-spec gas on a repeat basis."
)

# Escalation letter routing: administrative chain used when the AI-generated report is
# rendered as a formal escalation letter (From / To / Reference / Subject). Mirrors the
# discussion with the regulator: an Analyst-level query works up through Manager and
# Director to the Director-General, who signs on to the Ag. Director-General for matters
# requiring administrative action, and the Ag. Director-General in turn writes to the
# Authority Chief Executive for the final administrative decision.
ESCALATION_LETTER_ROUTES: dict[str, dict[str, str]] = {
    "ANALYST": {"from": "Analyst, Compliance Monitoring", "to": "Manager, Compliance Monitoring"},
    "MANAGER": {"from": "Manager, Compliance Monitoring", "to": "Director, Compliance Monitoring"},
    "DIRECTOR": {"from": "Director, Compliance Monitoring", "to": "Director-General"},
    "DG": {"from": "Director-General", "to": "Ag. Director-General"},
    "ACE": {"from": "Ag. Director-General", "to": "Authority Chief Executive"},
}

# Pilot SCADA integration actors. No live SCADA feed is wired up in this prototype; this list
# is surfaced in the UI (Settings) purely to document which operators are in scope for the
# pilot real-time data hookup discussed with the regulator, and is not a live connection.
SCADA_PILOT_ACTORS = [
    {"id": "NGIC", "name": "NNPC Gas Infrastructure Company", "status": "PILOT_PLANNED"},
    {"id": "SEPLAT", "name": "Seplat Energy", "status": "PILOT_PLANNED"},
]

# Reference "today" for the demo dataset (which spans 2026-01-01 to 2026-03-31). Used for the
# daily transporter population status (green/red) indicator and the 7-day rolling trend limit.
DEMO_TODAY = date(2026, 3, 31)

# Reference inputs for the theoretical line-pack estimate. Diameter/length are published
# figures for each pipeline; operating pressure is an assumed representative transmission
# pressure (no live SCADA pressure feed is wired up yet). Treat these as indicative,
# order-of-magnitude figures pending confirmation against actual pipe geometry and
# measured operating pressure from the network operator.
PIPELINE_REFERENCE_DATA = [
    {
        "id": "ELPS",
        "name": "Escravos-Lagos Pipeline System (ELPS)",
        "diameter_in": 36,
        "length_km": 342,
        "capacity_mmscfd": 2200,
        "assumed_operating_pressure_psig": 1000,
    },
    {
        "id": "AKK",
        "name": "Ajaokuta-Kaduna-Kano (AKK) Pipeline",
        "diameter_in": 40,
        "length_km": 614,
        "capacity_mmscfd": 3500,
        "assumed_operating_pressure_psig": 1000,
    },
    {
        "id": "OB3",
        "name": "OB3 / East-West Pipeline (Obiafu-Obrikom-Oben)",
        "diameter_in": 48,
        "length_km": 127,
        "capacity_mmscfd": 2400,
        "assumed_operating_pressure_psig": 1000,
    },
]


def _wrap_text(text: str, width: int) -> list[str]:
    words = text.split()
    lines: list[str] = []
    current = ""
    for word in words:
        candidate = f"{current} {word}".strip()
        if len(candidate) > width and current:
            lines.append(current)
            current = word
        else:
            current = candidate
    if current:
        lines.append(current)
    return lines or [""]


def slug(value: str) -> str:
    cleaned = re.sub(r"[^A-Za-z0-9]+", "-", value.strip().upper()).strip("-")
    return cleaned or "UNKNOWN"


def clean_text(value: Any) -> str:
    if value is None:
        return ""
    return re.sub(r"\s+", " ", str(value).replace("\n", " ")).strip()


def to_float(value: Any) -> float | None:
    if value is None or value == "":
        return None
    if isinstance(value, (int, float)):
        if isinstance(value, float) and math.isnan(value):
            return None
        return float(value)
    if isinstance(value, str):
        text = value.strip().replace(",", "")
        if not text or text.startswith("="):
            return None
        match = re.search(r"-?\d+(\.\d+)?", text)
        if not match:
            return None
        return float(match.group(0))
    return None


def to_iso(value: date | datetime | str | None) -> str | None:
    if value is None:
        return None
    if isinstance(value, datetime):
        return value.date().isoformat()
    if isinstance(value, date):
        return value.isoformat()
    return str(value)


def parse_date_cell(value: Any, previous: date | None = None) -> date | None:
    if isinstance(value, datetime):
        return value.date()
    if isinstance(value, date):
        return value
    if isinstance(value, str):
        text = value.strip()
        if not text:
            return None
        if text.startswith("=") and previous:
            match = re.match(r"=[A-Z]+\d+\+(\d+)", text)
            if match:
                return previous + timedelta(days=int(match.group(1)))
        for fmt in ("%Y-%m-%d", "%d-%b-%Y", "%d/%m/%Y", "%m/%d/%Y"):
            try:
                return datetime.strptime(text, fmt).date()
            except ValueError:
                pass
    return None


@dataclass
class Organisation:
    id: str
    code: str
    name: str
    roles: list[str]
    status: str = "ACTIVE"
    notes: str = ""


@dataclass
class Sector:
    id: str
    code: str
    name: str
    description: str


@dataclass
class TransporterProfile:
    organisation_id: str
    network_name: str
    operating_regions: list[str]
    license_ref: str = ""


@dataclass
class SupplierProfile:
    organisation_id: str
    gas_plant: str
    plant_capacity_mmscfd: float
    meter_streams: str
    meter_type: str
    dgdo_zone: str
    dgdo_target_mmscfd: float = 0.0


@dataclass
class CustomerProfile:
    organisation_id: str
    sector_id: str
    exit_point_id: str
    contract_volume_mmscfd: float
    customer_type: str


@dataclass
class EntryPoint:
    id: str
    code: str
    name: str
    location: str
    min_pressure_barg: float
    max_pressure_barg: float
    status: str = "ACTIVE"


@dataclass
class ExitPoint:
    id: str
    code: str
    name: str
    location: str
    min_pressure_barg: float
    max_pressure_barg: float
    shrinkage_threshold_mmscf: float
    status: str = "ACTIVE"


@dataclass
class NetworkEntryAgreement:
    id: str
    supplier_id: str
    entry_point_id: str
    effective_from: str
    effective_to: str
    min_pressure_barg: float
    max_pressure_barg: float
    status: str = "ACTIVE"


@dataclass
class TransporterSupplierAssignment:
    id: str
    transporter_id: str
    supplier_id: str
    period_id: str
    status: str = "ACTIVE"


@dataclass
class ReportingPeriod:
    id: str
    name: str
    start_date: str
    end_date: str
    status: str


@dataclass
class ThresholdRule:
    id: str
    scope: str
    metric: str
    warning_min: float | None
    warning_max: float | None
    severity: str


@dataclass
class TransporterSubmission:
    id: str
    transporter_id: str
    supplier_id: str
    period_id: str
    upload_source_type: str
    source_file_name: str
    status: str
    submitted_at: str
    validation_summary: dict[str, int]


@dataclass
class DailySupplierRecord:
    id: str
    submission_id: str
    date: str
    supplier_id: str
    entry_point_id: str
    entry_pressure_barg: float | None
    condensate_ltrs: float | None
    injection_mmscf: float | None
    mlf_mmscf: float | None
    effective_injection_mmscf: float | None
    attribution_mmscf: float
    shrinkage_mmscf: float
    shipper_imbalance_mmscf: float
    effective_attribution_mmscf: float
    supplier_imbalance_mmscf: float | None
    remark: str = ""


@dataclass
class DailyCustomerAttribution:
    id: str
    submission_id: str
    date: str
    supplier_id: str
    customer_id: str
    exit_point_id: str
    attribution_mmscf: float
    shrinkage_mmscf: float
    shipper_imbalance_mmscf: float
    effective_attribution_mmscf: float


@dataclass
class ValidationFlag:
    id: str
    submission_id: str
    record_type: str
    severity: str
    rule: str
    message: str
    affected_date: str | None = None
    supplier_id: str | None = None
    customer_id: str | None = None
    metric: str | None = None


@dataclass
class AIInsightEvidence:
    metric: str
    value: float | str
    label: str


@dataclass
class AIInsight:
    id: str
    submission_id: str
    scope: str
    severity: str
    title: str
    summary: str
    suggested_action: str
    date_range: str
    supplier_id: str | None = None
    customer_id: str | None = None
    sector_id: str | None = None
    evidence: list[AIInsightEvidence] = field(default_factory=list)


@dataclass
class EscalationNote:
    author: str
    stage: str
    action: str
    text: str
    timestamp: str


@dataclass
class GascoDailyNomination:
    """Daily nomination + quality submission from a gasco/supplier: previous day's volume,
    injection, a quality (lab) report reference, and today's projection — per the 'gascos
    provide information on previous day's volume and quality and injection, then put their
    projection for the day' workflow, backed by an uploaded gas quality report PDF."""

    id: str
    supplier_id: str
    date: str
    previous_day_volume_mmscf: float
    previous_day_injection_mmscf: float
    quality_report_file_name: str
    off_spec: bool
    projection_today_mmscf: float
    dispute_note: str
    submitted_at: str
    quality_report_path: str | None = None
    lab_mlf: float | None = None


@dataclass
class TransporterDirectEntry:
    """Direct daily data-entry record from a transporter, entered through the portal instead
    of an uploaded workbook — per 'rather than uploading a sheet, we give them ability to
    input'."""

    id: str
    transporter_id: str
    supplier_id: str
    date: str
    entry_pressure_barg: float | None
    condensate_ltrs: float | None
    injection_mmscf: float | None
    mlf_mmscf: float | None
    submitted_at: str


@dataclass
class ShipperNomination:
    """A shipper's own self-reported nominated volume for a given day, entered directly
    through the portal (with an optional supporting file) — independent of whatever
    figure the transporter's uploaded workbook shows for that shipper's linked supplier
    on the same day. Comparing the two side by side is what surfaces a discrepancy
    (e.g. transporter's data implies 10 MMScf, the shipper says it nominated 11)."""

    id: str
    shipper_id: str
    date: str
    nominated_volume_mmscf: float
    note: str
    submitted_at: str
    file_name: str | None
    stored_name: str | None


@dataclass
class EscalationCase:
    id: str
    flag_id: str
    rule: str
    severity: str
    message: str
    supplier_id: str | None
    customer_id: str | None
    affected_date: str | None
    stage: str
    status: str
    ai_report: str
    created_at: str
    updated_at: str
    notes: list[EscalationNote] = field(default_factory=list)


@dataclass
class KnowledgeDocument:
    """A policy, regulatory, or other reference document uploaded to the shared
    Knowledge Base — e.g. the PIA 2021, the NGTNC, NMDPRA circulars, internal SOPs,
    or submission templates. Kept separate from the Escalations/Cases workflow; this
    is a document library, not a workflow item."""

    id: str
    title: str
    category: str
    description: str
    tags: list[str]
    file_name: str | None
    stored_name: str | None
    mime_type: str | None
    size_bytes: int
    uploaded_by: str
    uploaded_at: str


@dataclass
class CaseNote:
    author: str
    action: str
    text: str
    timestamp: str


@dataclass
class ComplaintCase:
    """A manually logged complaint/case — e.g. a shipper disputing an attribution, a
    customer reporting a service issue, or a supplier raising a contractual query.
    Distinct from EscalationCase, which is opened from an AI-flagged exception; a
    ComplaintCase is opened directly by an analyst logging something reported to them."""

    id: str
    category: str
    subject: str
    description: str
    complainant_name: str
    complainant_org: str
    complainant_role: str
    complainant_email: str
    complainant_phone: str
    related_supplier_id: str | None
    related_customer_id: str | None
    related_transporter_id: str | None
    priority: str
    status: str
    assigned_to: str
    attachment_file_name: str | None
    attachment_path: str | None
    created_at: str
    updated_at: str
    notes: list[CaseNote] = field(default_factory=list)


class DataStore:
    def __init__(self) -> None:
        self.organisations: dict[str, Organisation] = {}
        self.sectors: dict[str, Sector] = {}
        self.transporters: dict[str, TransporterProfile] = {}
        self.shippers: dict[str, str] = {}
        self.shipper_supplier_map: dict[str, str] = {}
        self.suppliers: dict[str, SupplierProfile] = {}
        self.customers: dict[str, CustomerProfile] = {}
        self.entry_points: dict[str, EntryPoint] = {}
        self.exit_points: dict[str, ExitPoint] = {}
        self.neas: dict[str, NetworkEntryAgreement] = {}
        self.assignments: dict[str, TransporterSupplierAssignment] = {}
        self.periods: dict[str, ReportingPeriod] = {}
        self.thresholds: dict[str, ThresholdRule] = {}
        self.submissions: dict[str, TransporterSubmission] = {}
        self.supplier_records: list[DailySupplierRecord] = []
        self.customer_attributions: list[DailyCustomerAttribution] = []
        self.flags: list[ValidationFlag] = []
        self.insights: list[AIInsight] = []
        self.escalations: dict[str, EscalationCase] = {}
        self.gasco_nominations: list[GascoDailyNomination] = []
        self.direct_entries: list[TransporterDirectEntry] = []
        self.knowledge_docs: dict[str, KnowledgeDocument] = {}
        self.cases: dict[str, ComplaintCase] = {}
        self.shipper_nominations: list[ShipperNomination] = []
        self._counter = 0
        self.seed()

    def next_id(self, prefix: str) -> str:
        self._counter += 1
        return f"{prefix}-{self._counter:05d}"

    def seed(self) -> None:
        self.seed_sectors()
        self.seed_points()
        self.seed_organisations()
        self.seed_periods_thresholds_assignments()
        self.seed_demo_submission()
        self.seed_knowledge_base()
        self.seed_cases()
        self.seed_shipper_nominations()

    def add_org(self, code: str, name: str, roles: list[str], notes: str = "") -> str:
        org_id = slug(code)
        self.organisations[org_id] = Organisation(org_id, org_id, name, roles, notes=notes)
        return org_id

    def add_shipper_for_supplier(self, supplier_id: str) -> str:
        supplier = self.organisations[supplier_id]
        shipper_id = f"SHIPPER-{supplier_id}"
        self.organisations[shipper_id] = Organisation(
            shipper_id,
            shipper_id,
            supplier.name,
            ["SHIPPER"],
            notes=f"Mapped shipper account for {supplier.name}; linked to supplier/GASCO {supplier.code}.",
        )
        self.shippers[shipper_id] = supplier_id
        self.shipper_supplier_map[shipper_id] = supplier_id
        return shipper_id

    def shipper_id_for_supplier(self, supplier_id: str | None) -> str | None:
        if not supplier_id:
            return None
        return next((shipper_id for shipper_id, mapped_supplier in self.shipper_supplier_map.items() if mapped_supplier == supplier_id), None)

    def supplier_id_for_shipper(self, shipper_id: str | None) -> str | None:
        if not shipper_id:
            return None
        return self.shipper_supplier_map.get(shipper_id)

    def seed_sectors(self) -> None:
        for code, name, description in [
            ("GTP", "Gas to Power", "Power generation offtakers and grid-connected generation plants."),
            ("GBI", "Gas Based Industry", "Industrial customers using gas as production feedstock or process fuel."),
            ("GTC", "Gas to Commercial", "Commercial and distributed industrial gas customers."),
            ("LGD", "Local Gas Distributor", "Franchised distributors reticulating gas within a single city or industrial cluster."),
            ("RGD", "Regional Gas Distributor", "Distributors reticulating gas across multiple cities or states within a region."),
        ]:
            self.sectors[code] = Sector(code, code, name, description)

    def seed_points(self) -> None:
        for point in [
            EntryPoint("OBN-EP-01", "OBN-EP-01", "Oben Entry Point", "Edo / Oben gas hub", 55, 75),
            EntryPoint("ESC-EP-02", "ESC-EP-02", "Escravos Entry Point", "Western Niger Delta", 50, 70),
            EntryPoint("UQO-EP-03", "UQO-EP-03", "Uquo Entry Point", "Akwa Ibom / Uquo field", 45, 72),
            EntryPoint("UTG-EP-04", "UTG-EP-04", "Utorogu Entry Point", "Delta / Utorogu gas plant", 48, 74),
            EntryPoint("ITK-EP-05", "ITK-EP-05", "Itoki Interface", "Lagos / Itoki city gate", 35, 62),
            EntryPoint("SAP-EP-06", "SAP-EP-06", "Sapele Entry Point", "Delta / Sapele plant", 46, 70),
        ]:
            self.entry_points[point.id] = point

        for point in [
            ExitPoint("EGB-EX-01", "EGB-EX-01", "Egbin Delivery Point", "Ikorodu, Lagos", 30, 55, 0.8),
            ExitPoint("CAL-EX-02", "CAL-EX-02", "Calabar Delivery Point", "Cross River", 28, 52, 0.75),
            ExitPoint("IBM-EX-03", "IBM-EX-03", "Ibom Power Delivery Point", "Ikot Abasi", 28, 55, 0.7),
            ExitPoint("LAF-EX-04", "LAF-EX-04", "Mfamosing Cement Delivery Point", "Cross River", 25, 50, 0.6),
            ExitPoint("GLI-EX-05", "GLI-EX-05", "Greater Lagos Industrial Area", "Lagos industrial network", 22, 48, 0.55),
            ExitPoint("TAM-EX-06", "TAM-EX-06", "Trans Amadi Industrial Area", "Port Harcourt", 22, 48, 0.55),
            ExitPoint("SAG-EX-07", "SAG-EX-07", "Sagamu Industrial Corridor", "Ogun State", 22, 48, 0.55),
        ]:
            self.exit_points[point.id] = point

        for prefix in "ABCDEFGHIJ":
            for number in range(1, 6):
                exit_id = f"{prefix}{number}-EX"
                self.exit_points[exit_id] = ExitPoint(
                    exit_id,
                    exit_id,
                    f"Customer-{prefix}{number} Offtake Point",
                    "Demo network",
                    22,
                    55,
                    0.6,
                )

    def seed_organisations(self) -> None:
        self.add_org("NGIC", "NNPC Gas Infrastructure Company", ["TRANSPORTER"], "NNPC Gas & Power infrastructure operator.")
        self.add_org("NGML", "NNPC Gas Marketing Limited", ["SHIPPER", "GAS_MARKETER", "CUSTOMER"], "Gas marketing and local distribution market actor.")
        self.add_org("GASLINK", "Gaslink Nigeria Limited", ["GAS_MARKETER", "CUSTOMER"], "Greater Lagos gas distribution market actor.")
        self.add_org("CHGC", "Central Horizon Gas Company", ["GAS_MARKETER", "CUSTOMER"], "South-South gas distribution market actor.")
        self.add_org("GNSL", "Gas Network Services Limited", ["GAS_MARKETER"], "Virtual pipeline CNG distribution market actor.")
        self.add_org("TGNL", "Transit Gas Nigeria Limited", ["GAS_MARKETER"], "Sagamu/Iwopin gas distribution market actor.")
        self.add_org("SNG", "Shell Nigeria Gas", ["GAS_MARKETER", "CUSTOMER"], "Industrial and commercial gas distribution market actor.")
        self.add_org("ACCUGAS", "Accugas / Savannah Energy", ["SUPPLIER", "GAS_MARKETER"], "Uquo/Accugas gas processing network.")

        self.transporters["NGIC"] = TransporterProfile("NGIC", "National Gas Transmission Network", ["National", "Western", "Eastern"])

        supplier_rows = [
            ("SUPPLIER-A", "TotalEnergies", "Imo River Gas Plant", 150, "3", "Ultrasonic", "Engr. Adamu Kano", "OBN-EP-01", 88.5),
            ("SUPPLIER-B", "Supplier-B", "Escravos Gas Plant", 200, "2", "Turbine / Orifice", "Engr. Chioma Eze", "ESC-EP-02", 130.0),
            ("SEPLAT", "Seplat Energy", "Oben Gas Processing Plant", 525, "Multi-stream", "Ultrasonic", "DGDO West", "OBN-EP-01", 127.5),
            ("CNL", "Chevron Nigeria Limited", "Escravos Gas Plant", 680, "Multi-stream", "Ultrasonic", "DGDO West", "ESC-EP-02", 160.0),
            ("NEPL-OREDO", "NEPL (Oredo)", "Oredo Gas Facility", 360, "Multi-stream", "Orifice", "DGDO Mid-West", "UTG-EP-04", 210.0),
            ("SPDC", "SPDC / Renaissance Legacy", "Associated Gas Plants", 500, "Multi-stream", "Mixed", "DGDO East", "SAP-EP-06", 290.0),
            ("OML-34-JV", "OML-34 JV", "Utorogu Gas Plant", 360, "Multi-stream", "Ultrasonic", "DGDO West", "UTG-EP-04", 210.0),
            ("ACCUGAS", "Accugas / Savannah Energy", "Uquo Gas Processing Facility", 200, "Multi-stream", "Ultrasonic", "DGDO South-East", "UQO-EP-03", 115.0),
        ]
        for code, name, plant, capacity, streams, meter_type, dgdo_zone, entry_id, dgdo_target in supplier_rows:
            org_id = self.add_org(code, name, ["SUPPLIER"] if code != "ACCUGAS" else ["SUPPLIER", "GAS_MARKETER"])
            self.suppliers[org_id] = SupplierProfile(org_id, plant, capacity, streams, meter_type, dgdo_zone, dgdo_target)
            self.add_shipper_for_supplier(org_id)
            entry = self.entry_points[entry_id]
            nea_id = f"NEA-{org_id}"
            self.neas[nea_id] = NetworkEntryAgreement(
                nea_id,
                org_id,
                entry_id,
                "2026-01-01",
                "2026-12-31",
                entry.min_pressure_barg,
                entry.max_pressure_barg,
            )

        for letter in "CDEFGHIJ":
            org_id = self.add_org(f"SUPPLIER-{letter}", f"Supplier-{letter}", ["SUPPLIER"])
            self.suppliers[org_id] = SupplierProfile(org_id, f"Supplier-{letter} Gas Plant", 120, "2", "Ultrasonic", "DGDO Demo", 70.0)
            self.add_shipper_for_supplier(org_id)
            entry = "OBN-EP-01" if letter in "CEGI" else "ESC-EP-02"
            self.neas[f"NEA-{org_id}"] = NetworkEntryAgreement(f"NEA-{org_id}", org_id, entry, "2026-01-01", "2026-12-31", 50, 75)

        customer_rows = [
            ("EGBIN", "Egbin Power", "GTP", "EGB-EX-01", 220),
            ("CALABAR-GEN", "Calabar Generation Company", "GTP", "CAL-EX-02", 160),
            ("IBOM-POWER", "Ibom Power Company", "GTP", "IBM-EX-03", 120),
            ("LAFARGE", "Lafarge Africa", "GBI", "LAF-EX-04", 70),
            ("NOTORE", "Notore Chemical Industries", "GBI", "TAM-EX-06", 45),
            ("CADBURY", "Cadbury Nigeria", "GTC", "GLI-EX-05", 12),
            ("FLOUR-MILLS", "Flour Mills of Nigeria", "GBI", "GLI-EX-05", 24),
            ("DANGOTE-SUGAR", "Dangote Sugar", "GBI", "GLI-EX-05", 28),
            ("NIGERIAN-BREWERIES", "Nigerian Breweries", "GBI", "SAG-EX-07", 18),
            ("RITE-FOODS", "Rite Foods", "GTC", "SAG-EX-07", 10),
            ("COLEMAN-CABLES", "Coleman Cables", "GBI", "SAG-EX-07", 9),
            ("IKEJA-CITYGATE-LGD", "Ikeja City Gate Distribution Co.", "LGD", "GLI-EX-05", 15),
            ("PH-CITYGATE-LGD", "Port Harcourt City Gate Distribution Co.", "LGD", "TAM-EX-06", 12),
            ("SOUTHWEST-RGD", "South-West Regional Gas Distribution Co.", "RGD", "SAG-EX-07", 40),
            ("SOUTHSOUTH-RGD", "South-South Regional Gas Distribution Co.", "RGD", "CAL-EX-02", 35),
        ]
        for code, name, sector, exit_id, contract in customer_rows:
            org_id = self.add_org(code, name, ["CUSTOMER"])
            self.customers[org_id] = CustomerProfile(org_id, sector, exit_id, contract, self.sectors[sector].name)

        sector_cycle = ["GTP"] * 16 + ["GTC"] * 16 + ["GBI"] * 10 + ["LGD"] * 5 + ["RGD"] * 3
        index = 0
        for prefix in "ABCDEFGHIJ":
            for number in range(1, 6):
                code = f"CUSTOMER-{prefix}{number}"
                org_id = self.add_org(code, f"Customer-{prefix}{number}", ["CUSTOMER"])
                sector = sector_cycle[index]
                self.customers[org_id] = CustomerProfile(org_id, sector, f"{prefix}{number}-EX", 30 + number * 2, self.sectors[sector].name)
                index += 1

    def seed_periods_thresholds_assignments(self) -> None:
        self.periods["2026"] = ReportingPeriod("2026", "FY 2026", "2026-01-01", "2026-12-31", "OPEN")
        self.periods["JAN-2026"] = ReportingPeriod("JAN-2026", "January 2026", "2026-01-01", "2026-01-31", "OPEN")
        self.periods["FEB-2026"] = ReportingPeriod("FEB-2026", "February 2026", "2026-02-01", "2026-02-28", "OPEN")
        self.periods["MAR-2026"] = ReportingPeriod("MAR-2026", "March 2026", "2026-03-01", "2026-03-31", "OPEN")
        self.periods["Q1-2026"] = ReportingPeriod("Q1-2026", "Q1 2026", "2026-01-01", "2026-03-31", "OPEN")

        for metric, minimum, maximum, severity in [
            ("PRESSURE_BARG", None, None, "WARNING"),
            ("SHRINKAGE_MMSCF", None, 1.0, "WARNING"),
            ("SHIPPER_IMBALANCE_MMSCF", -2.0, 2.0, "WARNING"),
            ("CONDENSATE_LTRS", None, 350, "WARNING"),
            ("MISSING_DATA", None, None, "ERROR"),
            ("NOMINATION_VARIANCE_PCT", -NOMINATION_VARIANCE_THRESHOLD_PCT, NOMINATION_VARIANCE_THRESHOLD_PCT, "WARNING"),
            ("LINE_PACK_VARIANCE_PCT", -LINE_PACK_VARIANCE_THRESHOLD_PCT, LINE_PACK_VARIANCE_THRESHOLD_PCT, "WARNING"),
            ("MLF_LOSS_PCT", None, MLF_LOSS_THRESHOLD_PCT, "WARNING"),
        ]:
            rule_id = f"RULE-{metric}"
            self.thresholds[rule_id] = ThresholdRule(rule_id, "GLOBAL", metric, minimum, maximum, severity)

        for supplier_id in self.suppliers:
            self.assignments[f"ASN-NGIC-{supplier_id}"] = TransporterSupplierAssignment(
                f"ASN-NGIC-{supplier_id}", "NGIC", supplier_id, "2026"
            )

    def seed_demo_submission(self) -> None:
        # Deterministic 3-month seed (Jan 1 - Mar 31 2026) for three suppliers.
        # No random module: index-driven sine/cosine math keeps data stable across restarts.
        supplier_configs = [
            {
                "supplier_id": "SUPPLIER-A",
                "entry_point_id": "OBN-EP-01",
                "customers": ["CUSTOMER-A1", "CUSTOMER-B2", "CUSTOMER-D1", "CUSTOMER-I1", "CUSTOMER-J4"],
                "inj_base": 88.5, "inj_amp": 3.5,
                "nomination_mmscfd": 88.5,
                "mlf_base": 0.99775, "mlf_amp": 0.00075,
                "press_base": 64.0, "press_amp": 4.0,
                "cond_base": 240.0, "cond_amp": 40.0,
                "press_breach_days": {12, 31, 47, 63, 80},
                "cond_spike_days": {19, 55, 74},
                "mlf_spike_days": {27, 66},
                "cust_share": [0.36, 0.24, 0.20, 0.11, 0.09],
            },
            {
                "supplier_id": "SEPLAT",
                "entry_point_id": "OBN-EP-01",
                "customers": ["CUSTOMER-A2", "CUSTOMER-C4", "CUSTOMER-D2", "CUSTOMER-I2", "CUSTOMER-J5"],
                "inj_base": 127.5, "inj_amp": 7.5,
                "nomination_mmscfd": 127.5,
                "mlf_base": 0.9985, "mlf_amp": 0.0005,
                "press_base": 65.0, "press_amp": 7.0,
                "cond_base": 220.0, "cond_amp": 40.0,
                "press_breach_days": set(),
                "cond_spike_days": set(),
                "mlf_spike_days": {41},
                "cust_share": [0.34, 0.25, 0.22, 0.11, 0.08],
            },
            {
                "supplier_id": "CNL",
                "entry_point_id": "ESC-EP-02",
                "customers": ["CUSTOMER-A3", "CUSTOMER-C5", "CUSTOMER-F2", "CUSTOMER-I3", "CUSTOMER-J1"],
                "inj_base": 160.0, "inj_amp": 10.0,
                "nomination_mmscfd": 160.0,
                "mlf_base": 0.9980, "mlf_amp": 0.0005,
                "press_base": 61.5, "press_amp": 6.5,
                "cond_base": 265.0, "cond_amp": 45.0,
                "press_breach_days": set(),
                "cond_spike_days": set(),
                "mlf_spike_days": {58, 85},
                "cust_share": [0.32, 0.27, 0.24, 0.10, 0.07],
            },
        ]

        start = date(2026, 1, 1)
        end = date(2026, 3, 31)
        total_days = (end - start).days + 1

        for cfg in supplier_configs:
            supplier_id = cfg["supplier_id"]
            submission_id = f"SUB-DEMO-{supplier_id}"
            warning_count = 0
            customers = cfg["customers"]
            shares = cfg["cust_share"]

            for i in range(total_days):
                day = (start + timedelta(days=i)).isoformat()

                injection = round(cfg["inj_base"] + math.sin(i * 0.31) * cfg["inj_amp"] + math.cos(i * 0.13) * (cfg["inj_amp"] * 0.4), 2)
                mlf = round(cfg["mlf_base"] + math.sin(i * 0.21 + 1.1) * cfg["mlf_amp"], 5)
                if i in cfg["mlf_spike_days"]:
                    mlf = round(cfg["mlf_base"] - cfg["mlf_amp"] * 9, 5)

                pressure = round(cfg["press_base"] + math.sin(i * 0.27 + 0.6) * cfg["press_amp"], 1)
                if i in cfg["press_breach_days"]:
                    pressure = round(76.5 + abs(math.sin(i * 0.9)) * 3.5, 1)

                condensate = round(cfg["cond_base"] + math.sin(i * 0.18 + 0.4) * cfg["cond_amp"], 0)
                if i in cfg["cond_spike_days"]:
                    condensate = round(360 + abs(math.cos(i * 0.7)) * 35, 0)

                effective = injection * mlf
                # Attribution sums to ~95-98% of effective injection.
                attr_factor = 0.955 + (math.sin(i * 0.23 + 0.9) + 1) * 0.0125  # 0.955..0.980
                total_attr = effective * attr_factor

                attributions: list[float] = []
                shrinkages: list[float] = []
                imbalances: list[float] = []
                for ci, share in enumerate(shares):
                    attribution = round(total_attr * share, 3)
                    shrinkage = round(0.3 + (math.sin(i * 0.29 + ci * 1.3) + 1) * 0.1, 3)  # 0.3..0.5
                    imbalance = round(math.sin(i * 0.37 + ci * 0.7) * 0.1, 3)  # +/-0.1
                    attributions.append(attribution)
                    shrinkages.append(shrinkage)
                    imbalances.append(imbalance)

                day_total_attr = round(sum(attributions), 3)
                day_total_shrinkage = round(sum(shrinkages), 3)
                day_total_imbalance = round(sum(imbalances), 3)
                day_total_effective_attr = round(sum(a - s - im for a, s, im in zip(attributions, shrinkages, imbalances)), 3)
                supplier_imbalance = round(effective - day_total_attr, 3)

                self.supplier_records.append(
                    DailySupplierRecord(
                        f"DSR-DEMO-{supplier_id}-{i}",
                        submission_id,
                        day,
                        supplier_id,
                        cfg["entry_point_id"],
                        pressure,
                        condensate,
                        injection,
                        mlf,
                        round(effective, 3),
                        day_total_attr,
                        day_total_shrinkage,
                        day_total_imbalance,
                        day_total_effective_attr,
                        supplier_imbalance,
                        "",
                    )
                )

                for customer_id, attribution, shrinkage, imbalance in zip(customers, attributions, shrinkages, imbalances):
                    customer = self.customers[customer_id]
                    self.customer_attributions.append(
                        DailyCustomerAttribution(
                            f"DCA-DEMO-{supplier_id}-{i}-{customer_id}",
                            submission_id,
                            day,
                            supplier_id,
                            customer_id,
                            customer.exit_point_id,
                            attribution,
                            shrinkage,
                            imbalance,
                            round(attribution - shrinkage - imbalance, 3),
                        )
                    )

                if pressure > 75:
                    warning_count += 1
                    self.flags.append(
                        ValidationFlag(
                            f"FLAG-DEMO-{supplier_id}-{i}-PRESSURE",
                            submission_id,
                            "SUPPLIER_DAY",
                            "WARNING",
                            "PRESSURE_OUTSIDE_NEA",
                            f"Entry pressure {pressure:.1f} Barg is above the NEA maximum of 75 Barg.",
                            day,
                            supplier_id,
                            metric="PRESSURE_BARG",
                        )
                    )
                if condensate > 350:
                    warning_count += 1
                    self.flags.append(
                        ValidationFlag(
                            f"FLAG-DEMO-{supplier_id}-{i}-CONDENSATE",
                            submission_id,
                            "SUPPLIER_DAY",
                            "WARNING",
                            "CONDENSATE_THRESHOLD",
                            f"Condensate drop out {condensate:.0f} Ltrs is above the configured threshold.",
                            day,
                            supplier_id,
                            metric="CONDENSATE_LTRS",
                        )
                    )

                nomination = cfg["nomination_mmscfd"]
                nomination_variance_pct = round(((injection - nomination) / nomination) * 100, 2) if nomination else 0.0
                if abs(nomination_variance_pct) > NOMINATION_VARIANCE_THRESHOLD_PCT:
                    warning_count += 1
                    direction = "over-delivery" if nomination_variance_pct > 0 else "under-delivery"
                    self.flags.append(
                        ValidationFlag(
                            f"FLAG-DEMO-{supplier_id}-{i}-NOMINATION",
                            submission_id,
                            "SUPPLIER_DAY",
                            "WARNING",
                            "NOMINATION_VARIANCE",
                            f"Daily injection {injection:.1f} MMScf is {abs(nomination_variance_pct):.1f}% {direction} against the {nomination:.1f} MMScf nomination (band +/-{NOMINATION_VARIANCE_THRESHOLD_PCT:.0f}%).",
                            day,
                            supplier_id,
                            metric="NOMINATION_VARIANCE_PCT",
                        )
                    )

                line_pack_variance_pct = round(((injection - day_total_effective_attr) / injection) * 100, 2) if injection else 0.0
                if abs(line_pack_variance_pct) > LINE_PACK_VARIANCE_THRESHOLD_PCT:
                    warning_count += 1
                    self.flags.append(
                        ValidationFlag(
                            f"FLAG-DEMO-{supplier_id}-{i}-LINEPACK",
                            submission_id,
                            "SUPPLIER_DAY",
                            "WARNING",
                            "LINE_PACK_VARIANCE",
                            f"Injection vs. uptake gap is {abs(line_pack_variance_pct):.1f}% (injection {injection:.1f} MMScf vs. effective uptake {day_total_effective_attr:.1f} MMScf), above the {LINE_PACK_VARIANCE_THRESHOLD_PCT:.0f}% agreed threshold.",
                            day,
                            supplier_id,
                            metric="LINE_PACK_VARIANCE_PCT",
                        )
                    )

                mlf_loss_pct = round((1 - mlf) * 100, 3)
                if mlf_loss_pct > MLF_LOSS_THRESHOLD_PCT:
                    warning_count += 1
                    self.flags.append(
                        ValidationFlag(
                            f"FLAG-DEMO-{supplier_id}-{i}-MLF",
                            submission_id,
                            "SUPPLIER_DAY",
                            "WARNING",
                            "MLF_LOSS_THRESHOLD",
                            f"Measurement loss factor implies a {mlf_loss_pct:.2f}% loss on injection, above the {MLF_LOSS_THRESHOLD_PCT:.1f}% agreed band — worth a meter linearity check.",
                            day,
                            supplier_id,
                            metric="MLF_LOSS_PCT",
                        )
                    )

            submission = TransporterSubmission(
                submission_id,
                "NGIC",
                supplier_id,
                "Q1-2026",
                "TRANSPORTER",
                f"{supplier_id.lower()}-q1-2026.xlsx",
                "ACCEPTED_WITH_WARNINGS",
                datetime.now(UTC).isoformat(timespec="seconds").replace("+00:00", "Z"),
                {"errors": 0, "warnings": warning_count, "acceptedRecords": total_days},
            )
            self.submissions[submission_id] = submission
            self.generate_ai_insights(submission_id)

    def find_org_by_name_or_code(self, value: str, role: str | None = None) -> str | None:
        target = slug(value)
        aliases = {
            "SUPPLIER-A-INDOMIE": "SUPPLIER-A",
            "SUPPLIER-A-INDOMIE-GAS-LTD": "SUPPLIER-A",
            "EGBIN-POWER": "EGBIN",
            "NEPL-OREDO": "NEPL-OREDO",
            "OML-34-JV": "OML-34-JV",
            "SNG-WEST": "SNG",
            "NGML-WEST": "NGML",
        }
        if target in aliases:
            target = aliases[target]
        if target in self.organisations:
            org = self.organisations[target]
            if role is None or role in org.roles:
                return org.id
        for org in self.organisations.values():
            if slug(org.name) == target or target in slug(org.name):
                if role is None or role in org.roles:
                    return org.id
        return None

    def supplier_nea(self, supplier_id: str) -> NetworkEntryAgreement | None:
        for nea in self.neas.values():
            if nea.supplier_id == supplier_id and nea.status == "ACTIVE":
                return nea
        return None

    def is_assigned(self, transporter_id: str, supplier_id: str, period_id: str) -> bool:
        for assignment in self.assignments.values():
            if assignment.transporter_id == transporter_id and assignment.supplier_id == supplier_id and assignment.status == "ACTIVE":
                return assignment.period_id == period_id or assignment.period_id == "2026"
        return False

    def parse_workbook(self, file_bytes: bytes, file_name: str, transporter_id: str, supplier_id: str, period_id: str) -> dict[str, Any]:
        errors: list[ValidationFlag] = []
        warnings: list[ValidationFlag] = []

        # Load workbook first so we can auto-detect supplier and period
        try:
            workbook_values = load_workbook(BytesIO(file_bytes), data_only=True)
            workbook_formulas = load_workbook(BytesIO(file_bytes), data_only=False)
        except Exception as exc:
            errors.append(self.make_flag("", "SUBMISSION", "ERROR", "INVALID_WORKBOOK", f"Workbook could not be opened: {exc}"))
            return self.upload_response(None, errors, warnings)

        # Auto-detect supplier from sheet names if not provided
        if not supplier_id:
            for sid in self.suppliers:
                if self.pick_supplier_sheet(workbook_values.sheetnames, sid):
                    supplier_id = sid
                    break

        # Default period to full year if not provided
        if not period_id:
            period_id = "2026"

        # Default transporter
        if not transporter_id:
            transporter_id = "NGIC"

        if transporter_id not in self.transporters:
            errors.append(self.make_flag("", "SUBMISSION", "ERROR", "UNKNOWN_TRANSPORTER", "Selected transporter does not exist."))
        if supplier_id not in self.suppliers:
            errors.append(self.make_flag("", "SUBMISSION", "ERROR", "UNKNOWN_SUPPLIER", "Could not identify supplier from workbook. No matching sheet name found."))
        if period_id not in self.periods:
            errors.append(self.make_flag("", "SUBMISSION", "ERROR", "UNKNOWN_PERIOD", "Selected reporting period does not exist."))
        if errors:
            return self.upload_response(None, errors, warnings)

        if not self.is_assigned(transporter_id, supplier_id, period_id):
            errors.append(
                self.make_flag(
                    "",
                    "SUBMISSION",
                    "ERROR",
                    "TRANSPORTER_NOT_ASSIGNED",
                    "This transporter is not assigned to upload for the selected supplier.",
                    supplier_id=supplier_id,
                )
            )

        for submission in self.submissions.values():
            if submission.transporter_id == transporter_id and submission.supplier_id == supplier_id and submission.period_id == period_id:
                if submission.status in {"APPROVED", "LOCKED"}:
                    errors.append(
                        self.make_flag("", "SUBMISSION", "ERROR", "PERIOD_LOCKED", "Approved or locked submissions cannot be overwritten.")
                    )

        sheet_name = self.pick_supplier_sheet(workbook_values.sheetnames, supplier_id)
        if not sheet_name:
            errors.append(
                self.make_flag(
                    "",
                    "SUBMISSION",
                    "ERROR",
                    "SUPPLIER_SHEET_NOT_FOUND",
                    f"No sheet matched selected supplier {self.organisations[supplier_id].name}.",
                    supplier_id=supplier_id,
                )
            )
            return self.upload_response(None, errors, warnings)

        ws_values = workbook_values[sheet_name]
        ws_formulas = workbook_formulas[sheet_name]
        submission_id = self.next_id("SUB")

        metadata_supplier = clean_text(ws_values["A2"].value or ws_formulas["A2"].value)
        if metadata_supplier:
            matched_supplier = self.find_org_by_name_or_code(metadata_supplier, "SUPPLIER")
            if matched_supplier and matched_supplier != supplier_id:
                errors.append(
                    self.make_flag(
                        submission_id,
                        "SUBMISSION",
                        "ERROR",
                        "SUPPLIER_SELECTION_MISMATCH",
                        f"Workbook appears to contain {metadata_supplier}, not the selected supplier.",
                        supplier_id=supplier_id,
                    )
                )

        nea = self.supplier_nea(supplier_id)
        entry_point_id = nea.entry_point_id if nea else ""
        entry_point_text = self.metadata_value(ws_values, ws_formulas, "Entry Point:")
        if entry_point_text and nea:
            expected = self.entry_points[nea.entry_point_id]
            if expected.code not in entry_point_text and expected.name.upper() not in entry_point_text.upper():
                warnings.append(
                    self.make_flag(
                        submission_id,
                        "SUBMISSION",
                        "WARNING",
                        "ENTRY_POINT_MISMATCH",
                        f"Workbook entry point '{entry_point_text}' does not match configured NEA entry point {expected.name}.",
                        supplier_id=supplier_id,
                        metric="ENTRY_POINT",
                    )
                )
        elif nea:
            warnings.append(
                self.make_flag(
                    submission_id,
                    "SUBMISSION",
                    "WARNING",
                    "ENTRY_POINT_MISSING",
                    "Workbook entry point was blank; configured supplier NEA was used.",
                    supplier_id=supplier_id,
                    metric="ENTRY_POINT",
                )
            )

        period = self.periods[period_id]
        period_start = datetime.strptime(period.start_date, "%Y-%m-%d").date()
        period_end = datetime.strptime(period.end_date, "%Y-%m-%d").date()

        header_row, data_start_row = self.detect_layout(ws_values, ws_formulas)
        customer_groups = self.customer_groups(ws_values, ws_formulas, header_row)
        if not customer_groups:
            errors.append(self.make_flag(submission_id, "STRUCTURE", "ERROR", "CUSTOMER_COLUMNS_MISSING", "No grouped customer attribution columns were found."))

        supplier_records: list[DailySupplierRecord] = []
        customer_records: list[DailyCustomerAttribution] = []
        previous_day: date | None = None

        for row in range(data_start_row, ws_values.max_row + 1):
            raw_day = ws_values.cell(row, 1).value
            if raw_day is None:
                raw_day = ws_formulas.cell(row, 1).value
            day = parse_date_cell(raw_day, previous_day)
            if day is None:
                if row > 20:
                    break
                continue
            previous_day = day
            day_iso = day.isoformat()
            if not (period_start <= day <= period_end):
                errors.append(
                    self.make_flag(submission_id, "SUPPLIER_DAY", "ERROR", "DATE_OUTSIDE_PERIOD", f"{day_iso} is outside selected reporting period.", day_iso, supplier_id)
                )
                continue

            pressure = self.cell_number(ws_values, ws_formulas, row, 2)
            condensate = self.cell_number(ws_values, ws_formulas, row, 3)
            injection = self.cell_number(ws_values, ws_formulas, row, 4)
            mlf = self.cell_number(ws_values, ws_formulas, row, 5)
            effective = self.cell_number(ws_values, ws_formulas, row, 6)
            if effective is None and injection is not None and mlf is not None:
                effective = injection * mlf if mlf <= 2 else injection - mlf

            if injection is None:
                warnings.append(
                    self.make_flag(submission_id, "SUPPLIER_DAY", "WARNING", "MISSING_INJECTION", "Injection value is blank.", day_iso, supplier_id, metric="INJECTION_MMSCF")
                )

            if pressure is not None and nea:
                if pressure < nea.min_pressure_barg or pressure > nea.max_pressure_barg:
                    warnings.append(
                        self.make_flag(
                            submission_id,
                            "SUPPLIER_DAY",
                            "WARNING",
                            "PRESSURE_OUTSIDE_NEA",
                            f"Entry pressure {pressure:.2f} Barg is outside NEA band {nea.min_pressure_barg:.0f}-{nea.max_pressure_barg:.0f} Barg.",
                            day_iso,
                            supplier_id,
                            metric="PRESSURE_BARG",
                        )
                    )

            if condensate is not None and condensate > 350:
                warnings.append(
                    self.make_flag(
                        submission_id,
                        "SUPPLIER_DAY",
                        "WARNING",
                        "CONDENSATE_THRESHOLD",
                        f"Condensate drop out {condensate:.0f} Ltrs is above configured threshold.",
                        day_iso,
                        supplier_id,
                        metric="CONDENSATE_LTRS",
                    )
                )

            if mlf is not None and mlf <= 2:
                mlf_loss_pct = round((1 - mlf) * 100, 3)
                if mlf_loss_pct > MLF_LOSS_THRESHOLD_PCT:
                    warnings.append(
                        self.make_flag(
                            submission_id,
                            "SUPPLIER_DAY",
                            "WARNING",
                            "MLF_LOSS_THRESHOLD",
                            f"Measurement loss factor implies a {mlf_loss_pct:.2f}% loss on injection, above the {MLF_LOSS_THRESHOLD_PCT:.1f}% agreed band.",
                            day_iso,
                            supplier_id,
                            metric="MLF_LOSS_PCT",
                        )
                    )

            daily_customer_records: list[DailyCustomerAttribution] = []
            for group in customer_groups:
                customer_id = self.find_org_by_name_or_code(group["name"], "CUSTOMER")
                if not customer_id:
                    errors.append(
                        self.make_flag(
                            submission_id,
                            "CUSTOMER_DAY",
                            "ERROR",
                            "UNKNOWN_CUSTOMER",
                            f"Customer '{group['name']}' is not registered.",
                            day_iso,
                            supplier_id,
                            metric="CUSTOMER",
                        )
                    )
                    continue
                customer = self.customers[customer_id]
                col = group["col"]
                attribution = self.cell_number(ws_values, ws_formulas, row, col) or 0.0
                shrinkage = self.cell_number(ws_values, ws_formulas, row, col + 1) or 0.0
                imbalance = self.cell_number(ws_values, ws_formulas, row, col + 2) or 0.0
                effective_attr = self.cell_number(ws_values, ws_formulas, row, col + 3)
                if effective_attr is None:
                    effective_attr = attribution - shrinkage - imbalance
                daily_customer_records.append(
                    DailyCustomerAttribution(
                        self.next_id("DCA"),
                        submission_id,
                        day_iso,
                        supplier_id,
                        customer_id,
                        customer.exit_point_id,
                        attribution,
                        shrinkage,
                        imbalance,
                        effective_attr,
                    )
                )

                exit_point = self.exit_points[customer.exit_point_id]
                if abs(shrinkage) > exit_point.shrinkage_threshold_mmscf:
                    warnings.append(
                        self.make_flag(
                            submission_id,
                            "CUSTOMER_DAY",
                            "WARNING",
                            "SHRINKAGE_THRESHOLD",
                            f"{self.organisations[customer_id].name} shrinkage {shrinkage:.3f} MMScf is above exit point threshold.",
                            day_iso,
                            supplier_id,
                            customer_id,
                            "SHRINKAGE_MMSCF",
                        )
                    )

            total_attr = sum(item.attribution_mmscf for item in daily_customer_records)
            total_shrinkage = sum(item.shrinkage_mmscf for item in daily_customer_records)
            total_imbalance = sum(item.shipper_imbalance_mmscf for item in daily_customer_records)
            total_effective_attr = sum(item.effective_attribution_mmscf for item in daily_customer_records)
            supplier_imbalance = (effective - total_attr) if effective is not None else None
            if effective is not None and total_attr > effective:
                warnings.append(
                    self.make_flag(
                        submission_id,
                        "SUPPLIER_DAY",
                        "WARNING",
                        "ATTRIBUTION_EXCEEDS_EFFECTIVE_INJECTION",
                        f"Total customer attribution {total_attr:.3f} exceeds effective injection {effective:.3f}.",
                        day_iso,
                        supplier_id,
                        metric="ATTRIBUTION_MMSCF",
                    )
                )

            supplier_records.append(
                DailySupplierRecord(
                    self.next_id("DSR"),
                    submission_id,
                    day_iso,
                    supplier_id,
                    entry_point_id,
                    pressure,
                    condensate,
                    injection,
                    mlf,
                    effective,
                    total_attr,
                    total_shrinkage,
                    total_imbalance,
                    total_effective_attr,
                    supplier_imbalance,
                    self.row_remark(ws_values, ws_formulas, row),
                )
            )
            customer_records.extend(daily_customer_records)

        if errors:
            return self.upload_response(None, errors, warnings)

        status = "ACCEPTED_WITH_WARNINGS" if warnings else "ACCEPTED"
        submission = TransporterSubmission(
            submission_id,
            transporter_id,
            supplier_id,
            period_id,
            "TRANSPORTER",
            file_name,
            status,
            datetime.now(UTC).isoformat(timespec="seconds").replace("+00:00", "Z"),
            {"errors": 0, "warnings": len(warnings), "acceptedRecords": len(supplier_records)},
        )
        self.submissions[submission.id] = submission
        self.supplier_records.extend(supplier_records)
        self.customer_attributions.extend(customer_records)
        self.flags.extend(warnings)
        self.generate_ai_insights(submission.id)
        return self.upload_response(submission, [], warnings)

    def upload_response(self, submission: TransporterSubmission | None, errors: list[ValidationFlag], warnings: list[ValidationFlag]) -> dict[str, Any]:
        return {
            "ok": submission is not None and not errors,
            "submission": asdict(submission) if submission else None,
            "errors": [asdict(flag) for flag in errors],
            "warnings": [asdict(flag) for flag in warnings],
        }

    def make_flag(
        self,
        submission_id: str,
        record_type: str,
        severity: str,
        rule: str,
        message: str,
        affected_date: str | None = None,
        supplier_id: str | None = None,
        customer_id: str | None = None,
        metric: str | None = None,
    ) -> ValidationFlag:
        return ValidationFlag(self.next_id("FLAG"), submission_id, record_type, severity, rule, message, affected_date, supplier_id, customer_id, metric)

    def pick_supplier_sheet(self, sheet_names: list[str], supplier_id: str) -> str | None:
        supplier = self.organisations[supplier_id]
        candidates = {supplier_id, supplier.code, slug(supplier.name)}
        for sheet_name in sheet_names:
            sheet_slug = slug(sheet_name)
            if sheet_slug in candidates or any(candidate in sheet_slug for candidate in candidates):
                return sheet_name
        if supplier_id.startswith("SUPPLIER-") and supplier_id.title() in sheet_names:
            return supplier_id.title()
        return None

    def metadata_value(self, ws_values: Any, ws_formulas: Any, label: str) -> str:
        target = label.upper()
        for row in range(1, 8):
            for col in range(1, 12):
                value = clean_text(ws_values.cell(row, col).value or ws_formulas.cell(row, col).value)
                if value.upper() == target:
                    return clean_text(ws_values.cell(row, col + 1).value or ws_formulas.cell(row, col + 1).value)
        return ""

    def detect_layout(self, ws_values: Any, ws_formulas: Any) -> tuple[int, int]:
        for row in range(1, min(ws_values.max_row, 15) + 1):
            first_cell = clean_text(ws_values.cell(row, 1).value or ws_formulas.cell(row, 1).value).upper()
            if first_cell != "DATE":
                continue
            has_customer_group = False
            for col in range(2, ws_values.max_column + 1):
                value = clean_text(ws_values.cell(row, col).value or ws_formulas.cell(row, col).value)
                if value.upper() == "TOTAL" or self.find_org_by_name_or_code(value, "CUSTOMER") or "CUSTOMER" in value.upper():
                    has_customer_group = True
                    break
            if has_customer_group:
                return row, row + 2
        return 8, 10

    def customer_groups(self, ws_values: Any, ws_formulas: Any, header_row: int) -> list[dict[str, Any]]:
        groups = []
        for col in range(7, ws_values.max_column + 1):
            value = clean_text(ws_values.cell(header_row, col).value or ws_formulas.cell(header_row, col).value)
            if not value:
                continue
            if value.upper() == "TOTAL":
                break
            if "CUSTOMER" in value.upper() or self.find_org_by_name_or_code(value, "CUSTOMER"):
                groups.append({"name": value, "col": col})
        return groups

    def cell_number(self, ws_values: Any, ws_formulas: Any, row: int, col: int) -> float | None:
        value = to_float(ws_values.cell(row, col).value)
        if value is not None:
            return value
        return to_float(ws_formulas.cell(row, col).value)

    def row_remark(self, ws_values: Any, ws_formulas: Any, row: int) -> str:
        for col in range(ws_values.max_column, 1, -1):
            header = clean_text(ws_values.cell(row - 1, col).value or ws_formulas.cell(row - 1, col).value)
            if "REMARK" in header.upper():
                return clean_text(ws_values.cell(row, col).value or ws_formulas.cell(row, col).value)
        return ""

    def generate_ai_insights(self, submission_id: str) -> None:
        existing = [item for item in self.insights if item.submission_id != submission_id]
        self.insights = existing
        records = [record for record in self.supplier_records if record.submission_id == submission_id]
        attributions = [record for record in self.customer_attributions if record.submission_id == submission_id]
        flags = [flag for flag in self.flags if flag.submission_id == submission_id]
        if not records:
            return
        supplier_id = records[0].supplier_id
        supplier_name = self.organisations[supplier_id].name
        shipper_id = self.shipper_id_for_supplier(supplier_id)
        shipper_name = self.organisations[shipper_id].name if shipper_id else supplier_name
        dates = sorted(record.date for record in records)
        total_injection = sum(record.injection_mmscf or 0 for record in records)
        total_effective = sum(record.effective_injection_mmscf or 0 for record in records)
        total_attr = sum(record.attribution_mmscf for record in records)
        total_shrinkage = sum(record.shrinkage_mmscf for record in records)
        pressure_warnings = [flag for flag in flags if flag.rule == "PRESSURE_OUTSIDE_NEA"]
        severity = "HIGH" if any(flag.severity == "ERROR" for flag in flags) else ("MEDIUM" if flags else "LOW")
        summary = (
            f"{supplier_name} delivered {total_effective:.1f} MMScf effective injection for linked shipper {shipper_name} across {len(records)} gas days. "
            f"Offtaker attribution totalled {total_attr:.1f} MMScf with {total_shrinkage:.2f} MMScf shrinkage."
        )
        if pressure_warnings:
            summary += f" {len(pressure_warnings)} pressure exception was detected against the NEA band."
        self.insights.append(
            AIInsight(
                self.next_id("AI"),
                submission_id,
                "SUPPLIER",
                severity,
                f"{shipper_name} / {supplier_name} operating summary",
                summary,
                "Review flagged days first, then compare linked offtaker withdrawals against effective injection.",
                f"{dates[0]} to {dates[-1]}",
                supplier_id=supplier_id,
                evidence=[
                    AIInsightEvidence("effectiveInjection", round(total_effective, 3), "Effective Injection"),
                    AIInsightEvidence("attribution", round(total_attr, 3), "Offtaker Attribution"),
                    AIInsightEvidence("flags", len(flags), "Validation Flags"),
                ],
            )
        )

        by_customer: dict[str, list[DailyCustomerAttribution]] = {}
        for item in attributions:
            by_customer.setdefault(item.customer_id, []).append(item)
        for customer_id, rows in by_customer.items():
            total_customer_attr = sum(row.attribution_mmscf for row in rows)
            total_customer_shrinkage = sum(row.shrinkage_mmscf for row in rows)
            customer_name = self.organisations[customer_id].name
            customer_flags = [flag for flag in flags if flag.customer_id == customer_id]
            if customer_flags or total_customer_shrinkage > 1.0:
                self.insights.append(
                    AIInsight(
                        self.next_id("AI"),
                        submission_id,
                        "CUSTOMER",
                        "MEDIUM" if customer_flags else "LOW",
                        f"{customer_name} offtaker attribution note",
                        f"{customer_name} received {total_customer_attr:.1f} MMScf with {total_customer_shrinkage:.2f} MMScf shrinkage in the selected period.",
                        "Check exit point shrinkage and daily imbalance movement before closing review.",
                        f"{dates[0]} to {dates[-1]}",
                        supplier_id=supplier_id,
                        customer_id=customer_id,
                        sector_id=self.customers[customer_id].sector_id,
                        evidence=[
                            AIInsightEvidence("attribution", round(total_customer_attr, 3), "Attribution"),
                            AIInsightEvidence("shrinkage", round(total_customer_shrinkage, 3), "Shrinkage"),
                        ],
                    )
                )

    def build_report(self, filters: dict[str, str]) -> dict[str, Any]:
        supplier_records = list(self.supplier_records)
        customer_records = list(self.customer_attributions)
        flags = list(self.flags)
        insights = list(self.insights)

        start = filters.get("startDate")
        end = filters.get("endDate")
        supplier_id = filters.get("supplier")
        shipper_id = filters.get("shipper")
        if shipper_id and not supplier_id:
            supplier_id = self.supplier_id_for_shipper(shipper_id)
        customer_id = filters.get("customer")
        sector_id = filters.get("sector")
        transporter_id = filters.get("transporter")
        entry_point_id = filters.get("entryPoint")
        exit_point_id = filters.get("exitPoint")
        exception_type = filters.get("exceptionType")
        severity = filters.get("severity")

        def submission_matches(submission_id: str) -> bool:
            submission = self.submissions.get(submission_id)
            if not submission:
                return False
            if transporter_id and submission.transporter_id != transporter_id:
                return False
            return True

        def date_matches(day: str) -> bool:
            if start and day < start:
                return False
            if end and day > end:
                return False
            return True

        customer_records = [
            item
            for item in customer_records
            if submission_matches(item.submission_id)
            and date_matches(item.date)
            and (not supplier_id or item.supplier_id == supplier_id)
            and (not customer_id or item.customer_id == customer_id)
            and (not sector_id or self.customers[item.customer_id].sector_id == sector_id)
            and (not exit_point_id or item.exit_point_id == exit_point_id)
        ]

        matching_supplier_keys = {(item.submission_id, item.supplier_id, item.date) for item in customer_records}
        customer_limited = bool(customer_id or sector_id or exit_point_id)
        supplier_records = [
            item
            for item in supplier_records
            if submission_matches(item.submission_id)
            and date_matches(item.date)
            and (not supplier_id or item.supplier_id == supplier_id)
            and (not entry_point_id or item.entry_point_id == entry_point_id)
            and (not customer_limited or (item.submission_id, item.supplier_id, item.date) in matching_supplier_keys)
        ]

        flags = [
            flag
            for flag in flags
            if (not flag.affected_date or date_matches(flag.affected_date))
            and (not supplier_id or flag.supplier_id == supplier_id)
            and (not customer_id or flag.customer_id == customer_id)
            and (not exception_type or flag.rule == exception_type)
            and (not severity or flag.severity == severity)
            and (not transporter_id or self.submissions.get(flag.submission_id, TransporterSubmission("", "", "", "", "", "", "", "", {})).transporter_id == transporter_id)
        ]

        allowed_supplier_days = {(record.submission_id, record.supplier_id, record.date) for record in supplier_records}
        allowed_customer_days = {
            (record.submission_id, record.supplier_id, record.customer_id, record.date) for record in customer_records
        }
        flags = [
            flag
            for flag in flags
            if (
                flag.customer_id
                and (flag.submission_id, flag.supplier_id, flag.customer_id, flag.affected_date) in allowed_customer_days
            )
            or (
                not flag.customer_id
                and (flag.submission_id, flag.supplier_id, flag.affected_date) in allowed_supplier_days
            )
        ]

        if exception_type or severity:
            flagged_keys = {(flag.submission_id, flag.supplier_id, flag.affected_date) for flag in flags}
            supplier_records = [
                record for record in supplier_records if (record.submission_id, record.supplier_id, record.date) in flagged_keys
            ]
            customer_records = [
                record for record in customer_records if (record.submission_id, record.supplier_id, record.date) in flagged_keys
            ]

        insights = [
            insight
            for insight in insights
            if (not supplier_id or insight.supplier_id == supplier_id)
            and (not customer_id or insight.customer_id == customer_id)
            and (not sector_id or insight.sector_id == sector_id)
            and (not severity or insight.severity == severity or severity == "WARNING")
            and (not transporter_id or self.submissions.get(insight.submission_id, TransporterSubmission("", "", "", "", "", "", "", "", {})).transporter_id == transporter_id)
        ]

        supplier_summary = self.supplier_summary(supplier_records, customer_records, insights)
        customer_summary = self.customer_summary(customer_records, insights)
        dgdr_dgdo = self.dgdr_dgdo_summary(supplier_records, customer_records)
        return {
            "filters": filters,
            "kpis": self.kpis(supplier_records, customer_records, flags),
            "dailySeries": self.daily_series(supplier_records, customer_records),
            "weeklyTrend": self.period_trend(supplier_records, customer_records, filters),
            "transportationBreakdown": self.transportation_breakdown(supplier_records, customer_records),
            "utilization": self.utilization_summary(customer_records, insights),
            "supplierSummary": supplier_summary,
            "gasSupplierSummary": self.gas_supplier_summary(supplier_summary, dgdr_dgdo),
            "shipperSummary": self.shipper_summary(supplier_summary),
            "customerSummary": customer_summary,
            "offtakerSummary": self.offtaker_summary(customer_summary),
            "sectorSummary": self.sector_summary(customer_records, insights),
            "dgdrDgdo": dgdr_dgdo,
            "networkLinePackIndicator": self.network_line_pack_indicator(supplier_records, customer_records),
            "linePackStatus": self.line_pack_status(supplier_records, customer_records),
            "pipelineLinePack": self.pipeline_line_pack_theoretical(),
            "exceptions": [self.flag_to_view(flag) for flag in flags],
            "aiInsights": [self.insight_to_view(item) for item in insights],
            "uploads": [self.submission_to_view(item) for item in self.submissions.values()],
        }

    def kpis(self, supplier_records: list[DailySupplierRecord], customer_records: list[DailyCustomerAttribution], flags: list[ValidationFlag]) -> dict[str, Any]:
        return {
            "totalInjection": round(sum(record.injection_mmscf or 0 for record in supplier_records), 2),
            "effectiveAttribution": round(sum(record.effective_attribution_mmscf for record in customer_records), 2),
            "activeSuppliers": len({record.supplier_id for record in supplier_records}),
            "activeCustomers": len({record.customer_id for record in customer_records}),
            "totalShrinkage": round(sum(record.shrinkage_mmscf for record in customer_records), 3),
            "totalImbalance": round(sum(record.shipper_imbalance_mmscf for record in customer_records), 3),
            "totalMlfLoss": round(sum((record.injection_mmscf or 0) - (record.effective_injection_mmscf or record.injection_mmscf or 0) for record in supplier_records), 3),
            "exceptions": len(flags),
            "warnings": len([flag for flag in flags if flag.severity == "WARNING"]),
            "errors": len([flag for flag in flags if flag.severity == "ERROR"]),
        }

    def daily_series(self, supplier_records: list[DailySupplierRecord], customer_records: list[DailyCustomerAttribution]) -> list[dict[str, Any]]:
        days = sorted({record.date for record in supplier_records} | {record.date for record in customer_records})
        return [
            {
                "date": day,
                "injection": round(sum(record.injection_mmscf or 0 for record in supplier_records if record.date == day), 2),
                "attribution": round(sum(record.attribution_mmscf for record in customer_records if record.date == day), 2),
                "effectiveAttribution": round(sum(record.effective_attribution_mmscf for record in customer_records if record.date == day), 2),
                "shrinkage": round(sum(record.shrinkage_mmscf for record in customer_records if record.date == day), 3),
                "mlfLoss": round(
                    sum(
                        (record.injection_mmscf or 0) - (record.effective_injection_mmscf or record.injection_mmscf or 0)
                        for record in supplier_records
                        if record.date == day
                    ),
                    3,
                ),
                "imbalance": round(sum(record.shipper_imbalance_mmscf for record in customer_records if record.date == day), 3),
            }
            for day in days
        ]

    def supplier_summary(
        self, supplier_records: list[DailySupplierRecord], customer_records: list[DailyCustomerAttribution], insights: list[AIInsight]
    ) -> list[dict[str, Any]]:
        rows = []
        for supplier_id in sorted({record.supplier_id for record in supplier_records}):
            records = [record for record in supplier_records if record.supplier_id == supplier_id]
            attr_rows = [record for record in customer_records if record.supplier_id == supplier_id]
            insight = next((item for item in insights if item.scope == "SUPPLIER" and item.supplier_id == supplier_id), None)
            rows.append(
                {
                    "supplierId": supplier_id,
                    "supplier": self.organisations[supplier_id].name,
                    "shipperId": self.shipper_id_for_supplier(supplier_id),
                    "shipper": self.organisations[self.shipper_id_for_supplier(supplier_id)].name if self.shipper_id_for_supplier(supplier_id) else self.organisations[supplier_id].name,
                    "days": len({record.date for record in records}),
                    "injection": round(sum(record.injection_mmscf or 0 for record in records), 2),
                    "effectiveInjection": round(sum(record.effective_injection_mmscf or 0 for record in records), 2),
                    "attribution": round(sum(record.attribution_mmscf for record in attr_rows), 2),
                    "shrinkage": round(sum(record.shrinkage_mmscf for record in attr_rows), 3),
                    "supplierImbalance": round(sum(record.supplier_imbalance_mmscf or 0 for record in records), 2),
                    "customers": len({record.customer_id for record in attr_rows}),
                    "offtakers": len({record.customer_id for record in attr_rows}),
                    "aiRemark": insight.summary if insight else "No AI exception detected for the active filters.",
                    "insightId": insight.id if insight else None,
                }
            )
        return rows

    def customer_summary(self, customer_records: list[DailyCustomerAttribution], insights: list[AIInsight]) -> list[dict[str, Any]]:
        rows = []
        for customer_id in sorted({record.customer_id for record in customer_records}):
            records = [record for record in customer_records if record.customer_id == customer_id]
            customer = self.customers[customer_id]
            insight = next((item for item in insights if item.scope == "CUSTOMER" and item.customer_id == customer_id), None)
            rows.append(
                {
                    "customerId": customer_id,
                    "customer": self.organisations[customer_id].name,
                    "offtakerId": customer_id,
                    "offtaker": self.organisations[customer_id].name,
                    "sector": customer.sector_id,
                    "supplierCount": len({record.supplier_id for record in records}),
                    "shipperCount": len({record.supplier_id for record in records}),
                    "exitPoint": self.exit_points[customer.exit_point_id].name,
                    "attribution": round(sum(record.attribution_mmscf for record in records), 2),
                    "shrinkage": round(sum(record.shrinkage_mmscf for record in records), 3),
                    "imbalance": round(sum(record.shipper_imbalance_mmscf for record in records), 3),
                    "effectiveAttribution": round(sum(record.effective_attribution_mmscf for record in records), 2),
                    "aiRemark": insight.summary if insight else "Customer activity is within expected pattern for the active filters.",
                    "insightId": insight.id if insight else None,
                }
            )
        return rows

    def shipper_summary(self, supplier_rows: list[dict[str, Any]]) -> list[dict[str, Any]]:
        rows = []
        for row in supplier_rows:
            supplier_id = row["supplierId"]
            shipper_id = self.shipper_id_for_supplier(supplier_id)
            days = max(row.get("days", 0), 1)
            profile = self.suppliers.get(supplier_id)
            nominated = round((profile.dgdo_target_mmscfd if profile else 0.0) * days, 1)
            delivered = row.get("injection", 0.0)
            uptake = row.get("effectiveAttribution", row.get("attribution", 0.0))
            rows.append(
                {
                    "shipperId": shipper_id,
                    "shipper": self.organisations[shipper_id].name if shipper_id else row["supplier"],
                    "linkedSupplierId": supplier_id,
                    "linkedSupplier": row["supplier"],
                    "days": row["days"],
                    "nominatedVolume": nominated,
                    "injection": delivered,
                    "effectiveOfftake": uptake,
                    "imbalance": row["supplierImbalance"],
                    "offtakers": row["offtakers"],
                    "variancePct": round(((delivered - nominated) / nominated) * 100, 1) if nominated else 0.0,
                    "aiRemark": row["aiRemark"],
                    "insightId": row["insightId"],
                }
            )
        return rows

    def offtaker_summary(self, customer_rows: list[dict[str, Any]]) -> list[dict[str, Any]]:
        return [
            {
                **row,
                "offtakerId": row["customerId"],
                "offtaker": row["customer"],
                "shipperCount": row["supplierCount"],
            }
            for row in customer_rows
        ]

    def gas_supplier_summary(self, supplier_rows: list[dict[str, Any]], dgdr_dgdo: dict[str, Any]) -> list[dict[str, Any]]:
        dgdo_by_supplier = {row["supplierId"]: row for row in dgdr_dgdo.get("bySupplier", [])}
        rows = []
        for row in supplier_rows:
            dgdo = dgdo_by_supplier.get(row["supplierId"], {})
            rows.append(
                {
                    **row,
                    "gasSupplierId": row["supplierId"],
                    "gasSupplier": row["supplier"],
                    "dgdoTarget": dgdo.get("dgdoTarget", 0),
                    "dgdoActual": dgdo.get("actual", row.get("injection", 0)),
                    "dgdoVariancePct": dgdo.get("variancePct", 0),
                }
            )
        return rows

    def sector_summary(self, customer_records: list[DailyCustomerAttribution], insights: list[AIInsight]) -> list[dict[str, Any]]:
        rows = []
        for sector_id in self.sectors:
            customer_ids = {cid for cid, profile in self.customers.items() if profile.sector_id == sector_id}
            records = [record for record in customer_records if record.customer_id in customer_ids]
            supplier_ids = sorted({record.supplier_id for record in records})
            rows.append(
                {
                    "sectorId": sector_id,
                    "sector": self.sectors[sector_id].name,
                    "customers": len({record.customer_id for record in records}),
                    "suppliers": len({record.supplier_id for record in records}),
                    "offtakerIds": sorted({record.customer_id for record in records}),
                    "shipperIds": [self.shipper_id_for_supplier(sid) for sid in supplier_ids if self.shipper_id_for_supplier(sid)],
                    "supplierIds": supplier_ids,
                    "attribution": round(sum(record.attribution_mmscf for record in records), 2),
                    "shrinkage": round(sum(record.shrinkage_mmscf for record in records), 3),
                    "effectiveAttribution": round(sum(record.effective_attribution_mmscf for record in records), 2),
                    "aiRemark": self.sector_remark(sector_id, records),
                }
            )
        return rows

    def sector_remark(self, sector_id: str, records: list[DailyCustomerAttribution]) -> str:
        if not records:
            return f"No {sector_id} activity for the active filters."
        total = sum(record.effective_attribution_mmscf for record in records)
        return f"{sector_id} accounts for {total:.1f} MMScf effective attribution across {len({r.customer_id for r in records})} offtakers."

    def dgdr_dgdo_summary(
        self, supplier_records: list[DailySupplierRecord], customer_records: list[DailyCustomerAttribution]
    ) -> dict[str, Any]:
        period_days = max(len({r.date for r in supplier_records} | {r.date for r in customer_records}), 1)

        by_sector = []
        overall_target = 0.0
        overall_actual = 0.0
        for sector_id in ["GTP", "GBI", "GTC"]:
            customer_ids = {cid for cid, profile in self.customers.items() if profile.sector_id == sector_id}
            records = [r for r in customer_records if r.customer_id in customer_ids]
            target = round(sum(self.customers[cid].contract_volume_mmscfd for cid in customer_ids) * period_days, 1)
            actual = round(sum(r.effective_attribution_mmscf for r in records), 1)
            overall_target += target
            overall_actual += actual
            by_sector.append(
                {
                    "sectorId": sector_id,
                    "sector": self.sectors[sector_id].name,
                    "dgdrTarget": target,
                    "actual": actual,
                    "variancePct": round(((actual - target) / target) * 100, 1) if target else 0.0,
                }
            )

        by_supplier = []
        target_shortfall_by_supplier = {"SUPPLIER-A": 0.18, "SEPLAT": 0.27, "CNL": 0.36}
        for supplier_id in sorted({r.supplier_id for r in supplier_records}):
            profile = self.suppliers.get(supplier_id)
            if not profile:
                continue
            actual = round(sum(r.injection_mmscf or 0 for r in supplier_records if r.supplier_id == supplier_id), 1)
            target = round(actual / (1 - target_shortfall_by_supplier.get(supplier_id, 0.22)), 1)
            by_supplier.append(
                {
                    "supplierId": supplier_id,
                    "supplier": self.organisations[supplier_id].name,
                    "dgdoTarget": target,
                    "actual": actual,
                    "variancePct": round(((actual - target) / target) * 100, 1) if target else 0.0,
                }
            )

        return {
            "periodDays": period_days,
            "overall": {
                "dgdrTarget": round(overall_target, 1),
                "actual": round(overall_actual, 1),
                "variancePct": round(((overall_actual - overall_target) / overall_target) * 100, 1) if overall_target else 0.0,
            },
            "bySector": by_sector,
            "bySupplier": by_supplier,
        }

    def period_trend(
        self,
        supplier_records: list[DailySupplierRecord],
        customer_records: list[DailyCustomerAttribution],
        filters: dict[str, str],
    ) -> dict[str, Any]:
        """Trend chart data that adapts its bucket size to the active time filter:
        a Week/Today filter (or no filter) buckets by day (last 7 days of data),
        a Month filter buckets by ISO week, and a Quarter/All filter buckets by month.
        This keeps the dashboard trend chart legible across the Today/Week/Month/Quarter
        presets instead of always showing a fixed last-7-days window."""
        all_dates = sorted({r.date for r in supplier_records} | {r.date for r in customer_records})
        if not all_dates:
            return {"granularity": "day", "points": []}

        start = filters.get("startDate") or all_dates[0]
        end = filters.get("endDate") or all_dates[-1]
        try:
            span_days = (date.fromisoformat(end) - date.fromisoformat(start)).days + 1
        except ValueError:
            span_days = len(all_dates)

        def injection_on(day: str) -> float:
            return sum(r.injection_mmscf or 0 for r in supplier_records if r.date == day)

        def uptake_on(day: str) -> float:
            return sum(r.effective_attribution_mmscf for r in customer_records if r.date == day)

        if span_days <= 9:
            days = all_dates[-7:]
            points = [{"label": d, "date": d, "injection": round(injection_on(d), 2), "uptake": round(uptake_on(d), 2)} for d in days]
            return {"granularity": "day", "points": points}

        if span_days <= 40:
            buckets: dict[str, list[str]] = {}
            for d in all_dates:
                iso = date.fromisoformat(d).isocalendar()
                key = f"{iso[0]}-W{iso[1]:02d}"
                buckets.setdefault(key, []).append(d)
            points = [
                {
                    "label": key,
                    "date": days_in[-1],
                    "injection": round(sum(injection_on(d) for d in days_in), 2),
                    "uptake": round(sum(uptake_on(d) for d in days_in), 2),
                }
                for key, days_in in sorted(buckets.items())[-8:]
            ]
            return {"granularity": "week", "points": points}

        buckets = {}
        for d in all_dates:
            key = d[:7]
            buckets.setdefault(key, []).append(d)
        points = [
            {
                "label": key,
                "date": days_in[-1],
                "injection": round(sum(injection_on(d) for d in days_in), 2),
                "uptake": round(sum(uptake_on(d) for d in days_in), 2),
            }
            for key, days_in in sorted(buckets.items())[-12:]
        ]
        return {"granularity": "month", "points": points}

    def transportation_breakdown(
        self, supplier_records: list[DailySupplierRecord], customer_records: list[DailyCustomerAttribution]
    ) -> dict[str, Any]:
        """Waterfall-style breakdown of where every injected molecule went: total
        injection carved into effective attribution, shrinkage, MLF (measurement) loss,
        and a residual imbalance/unaccounted bucket, so the pieces sum back to the whole."""
        total_injection = sum(r.injection_mmscf or 0 for r in supplier_records)
        total_effective_injection = sum(r.effective_injection_mmscf or (r.injection_mmscf or 0) for r in supplier_records)
        mlf_loss = max(total_injection - total_effective_injection, 0.0)
        effective_attribution = sum(r.effective_attribution_mmscf for r in customer_records)
        shrinkage = sum(r.shrinkage_mmscf for r in customer_records)
        imbalance = round(total_injection - mlf_loss - effective_attribution - shrinkage, 3)
        return {
            "totalInjection": round(total_injection, 2),
            "effectiveAttribution": round(effective_attribution, 2),
            "shrinkage": round(shrinkage, 3),
            "mlfLoss": round(mlf_loss, 3),
            "imbalance": imbalance,
        }

    def utilization_summary(self, customer_records: list[DailyCustomerAttribution], insights: list[AIInsight]) -> dict[str, Any]:
        """Utilization Performance view data: all-sector split, strategic-sector-only
        split, and the strategic sectors' combined actual vs. DGDR target."""
        all_sectors = self.sector_summary(customer_records, insights)
        strategic = [row for row in all_sectors if row["sectorId"] in ("GTP", "GBI", "GTC")]
        strategic_total = round(sum(row["effectiveAttribution"] for row in strategic), 2)
        distribution = [row for row in all_sectors if row["sectorId"] in ("LGD", "RGD")]
        return {
            "bySector": all_sectors,
            "byStrategicSector": strategic,
            "byDistributionSector": distribution,
            "strategicTotal": strategic_total,
        }

    def line_pack_status(
        self, supplier_records: list[DailySupplierRecord], customer_records: list[DailyCustomerAttribution]
    ) -> list[dict[str, Any]]:
        rows = []
        for supplier_id in sorted({r.supplier_id for r in supplier_records}):
            s_records = sorted((r for r in supplier_records if r.supplier_id == supplier_id), key=lambda r: r.date)
            if not s_records:
                continue
            latest = s_records[-1]
            uptake = sum(r.effective_attribution_mmscf for r in customer_records if r.supplier_id == supplier_id and r.date == latest.date)
            injection = latest.injection_mmscf or 0
            variance_pct = round(((injection - uptake) / injection) * 100, 2) if injection else 0.0
            rows.append(
                {
                    "supplierId": supplier_id,
                    "supplier": self.organisations[supplier_id].name,
                    "shipperId": self.shipper_id_for_supplier(supplier_id),
                    "shipper": self.organisations[self.shipper_id_for_supplier(supplier_id)].name if self.shipper_id_for_supplier(supplier_id) else self.organisations[supplier_id].name,
                    "date": latest.date,
                    "injection": round(injection, 2),
                    "uptake": round(uptake, 2),
                    "variancePct": variance_pct,
                    "breach": abs(variance_pct) > LINE_PACK_VARIANCE_THRESHOLD_PCT,
                }
            )
        return rows

    def network_line_pack_indicator(
        self, supplier_records: list[DailySupplierRecord], customer_records: list[DailyCustomerAttribution]
    ) -> dict[str, Any]:
        total_injection = sum(record.injection_mmscf or 0 for record in supplier_records)
        total_uptake = sum(record.effective_attribution_mmscf for record in customer_records)
        variance_pct = round(((total_injection - total_uptake) / total_injection) * 100, 2) if total_injection else 0.0
        abs_variance = abs(variance_pct)
        if abs_variance <= 3:
            status = "SAFE"
        elif abs_variance <= LINE_PACK_VARIANCE_THRESHOLD_PCT:
            status = "WARNING"
        else:
            status = "CRITICAL"
        return {
            "status": status,
            "variancePct": variance_pct,
            "totalInjection": round(total_injection, 2),
            "totalUptake": round(total_uptake, 2),
            "thresholdPct": LINE_PACK_VARIANCE_THRESHOLD_PCT,
        }

    def pipeline_line_pack_theoretical(self) -> list[dict[str, Any]]:
        base_pressure_psia = 14.65
        rows = []
        for pipe in PIPELINE_REFERENCE_DATA:
            diameter_ft = pipe["diameter_in"] / 12
            length_ft = pipe["length_km"] * 3280.84
            volume_ft3 = math.pi / 4 * diameter_ft**2 * length_ft
            operating_psia = pipe["assumed_operating_pressure_psig"] + 14.65
            line_pack_mmscf = volume_ft3 * (operating_psia / base_pressure_psia) / 1_000_000
            hours_of_capacity = (line_pack_mmscf / (pipe["capacity_mmscfd"] / 24)) if pipe["capacity_mmscfd"] else None
            rows.append(
                {
                    "id": pipe["id"],
                    "name": pipe["name"],
                    "diameterIn": pipe["diameter_in"],
                    "lengthKm": pipe["length_km"],
                    "capacityMmscfd": pipe["capacity_mmscfd"],
                    "assumedOperatingPressurePsig": pipe["assumed_operating_pressure_psig"],
                    "theoreticalLinePackMmscf": round(line_pack_mmscf, 1),
                    "hoursOfNameplateCapacity": round(hours_of_capacity, 1) if hours_of_capacity else None,
                }
            )
        return rows

    def recommended_action_for(self, rule: str) -> str:
        return RECOMMENDED_ACTIONS.get(rule, DEFAULT_RECOMMENDED_ACTION)

    def regulatory_reference_for(self, rule: str) -> dict[str, str]:
        return PIA_REFERENCES.get(rule, DEFAULT_PIA_REFERENCE)

    def flag_to_view(self, flag: ValidationFlag) -> dict[str, Any]:
        result = asdict(flag)
        if flag.supplier_id and flag.supplier_id in self.organisations:
            result["supplier"] = self.organisations[flag.supplier_id].name
            result["gasSupplier"] = self.organisations[flag.supplier_id].name
            result["shipperId"] = self.shipper_id_for_supplier(flag.supplier_id)
            result["shipper"] = self.organisations[result["shipperId"]].name if result["shipperId"] else self.organisations[flag.supplier_id].name
        if flag.customer_id and flag.customer_id in self.organisations:
            result["customer"] = self.organisations[flag.customer_id].name
            result["offtaker"] = self.organisations[flag.customer_id].name
        result["recommendedAction"] = self.recommended_action_for(flag.rule)
        result["regulatoryReference"] = self.regulatory_reference_for(flag.rule)
        escalation = next((case for case in self.escalations.values() if case.flag_id == flag.id), None)
        result["escalationId"] = escalation.id if escalation else None
        result["escalationStage"] = escalation.stage if escalation else None
        result["escalationStatus"] = escalation.status if escalation else None
        return result

    def insight_to_view(self, insight: AIInsight) -> dict[str, Any]:
        result = asdict(insight)
        if insight.supplier_id:
            result["supplier"] = self.organisations[insight.supplier_id].name
            result["gasSupplier"] = self.organisations[insight.supplier_id].name
            result["shipperId"] = self.shipper_id_for_supplier(insight.supplier_id)
            result["shipper"] = self.organisations[result["shipperId"]].name if result["shipperId"] else self.organisations[insight.supplier_id].name
        if insight.customer_id:
            result["customer"] = self.organisations[insight.customer_id].name
            result["offtaker"] = self.organisations[insight.customer_id].name
        if insight.sector_id:
            result["sector"] = self.sectors[insight.sector_id].name
        return result

    def submission_to_view(self, submission: TransporterSubmission) -> dict[str, Any]:
        return {
            **asdict(submission),
            "transporter": self.organisations.get(submission.transporter_id, Organisation("", "", "", [])).name,
            "supplier": self.organisations.get(submission.supplier_id, Organisation("", "", "", [])).name,
            "gasSupplier": self.organisations.get(submission.supplier_id, Organisation("", "", "", [])).name,
            "shipperId": self.shipper_id_for_supplier(submission.supplier_id),
            "shipper": self.organisations[self.shipper_id_for_supplier(submission.supplier_id)].name if self.shipper_id_for_supplier(submission.supplier_id) else self.organisations.get(submission.supplier_id, Organisation("", "", "", [])).name,
            "period": self.periods.get(submission.period_id, ReportingPeriod("", "", "", "", "")).name,
        }

    def update_threshold(self, rule_id: str, payload: dict[str, Any]) -> dict[str, Any]:
        if rule_id not in self.thresholds:
            return {"ok": False, "message": "Threshold rule not found."}
        rule = self.thresholds[rule_id]
        def _num(key: str) -> float | None:
            v = payload.get(key)
            if v is None or v == "":
                return None
            try:
                return float(v)
            except (TypeError, ValueError):
                return None
        if "warning_min" in payload:
            rule.warning_min = _num("warning_min")
        if "warning_max" in payload:
            rule.warning_max = _num("warning_max")
        if "severity" in payload and payload["severity"] in ("WARNING", "ERROR"):
            rule.severity = payload["severity"]
        return {"ok": True, "threshold": asdict(rule)}

    # ─── Escalation workflow ──────────────────────────────────────────────────
    def find_flag(self, flag_id: str) -> ValidationFlag | None:
        return next((f for f in self.flags if f.id == flag_id), None)

    def build_ai_report(self, flag: ValidationFlag) -> str:
        supplier_name = self.organisations[flag.supplier_id].name if flag.supplier_id and flag.supplier_id in self.organisations else "the supplier"
        customer_name = self.organisations[flag.customer_id].name if flag.customer_id and flag.customer_id in self.organisations else None
        ref = self.regulatory_reference_for(flag.rule)
        lines = [
            f"AI-generated potential incident report — {flag.rule.replace('_', ' ').title()}",
            f"Severity: {flag.severity}",
            f"Entity: {supplier_name}" + (f" / {customer_name}" if customer_name else ""),
            f"Date: {flag.affected_date or 'n/a'}",
            "",
            f"Finding: {flag.message}",
            "",
            f"Regulatory basis: {ref['act']}, {ref['section']}; {ref['code']}.",
            "",
            f"Recommended action (Consequence Management): {self.recommended_action_for(flag.rule)}",
        ]
        return "\n".join(lines)

    # ─── Escalation report as a formal letter (Word / PDF) ────────────────────
    def build_escalation_letter(self, case: EscalationCase) -> dict[str, Any]:
        view = self.escalation_to_view(case)
        route = ESCALATION_LETTER_ROUTES.get(case.stage, ESCALATION_LETTER_ROUTES["ANALYST"])
        ref = self.regulatory_reference_for(case.rule)
        entity = view.get("shipper") or view.get("supplier") or view.get("customer") or "the reporting entity"
        subject = f"Query on {case.rule.replace('_', ' ').title()} — {entity}" + (f" ({case.affected_date})" if case.affected_date else "")
        body = [
            f"This letter serves to formally notify you of a potential incident of non-conformity identified by the "
            f"Transporter Intelligence AI Consequence Management engine on the domestic gas transportation network.",
            f"Entity concerned: {entity}. Date of occurrence: {case.affected_date or 'n/a'}. Severity classification: {case.severity}.",
            f"Finding: {case.message}",
            f"Regulatory basis: {ref['act']}, {ref['section']}; {ref['code']}.",
            f"Recommended action: {self.recommended_action_for(case.rule)}",
            "You are kindly requested to review the above and revert with your comments/explanation, or confirm the "
            "corrective action taken, within five (5) working days of receipt of this letter.",
        ]
        approval_trail = [
            {
                "author": note.author,
                "stage": ESCALATION_STAGE_LABELS.get(note.stage, note.stage),
                "action": note.action,
                "timestamp": note.timestamp,
            }
            for note in case.notes
        ]
        comments = [note.text for note in case.notes if note.action == "COMMENT"]
        return {
            "referenceNumber": f"NMDPRA/TI/{case.id}",
            "date": to_iso(datetime.now(UTC)),
            "fromTitle": route["from"],
            "toTitle": route["to"],
            "subject": subject,
            "body": body,
            "approvalStage": ESCALATION_STAGE_LABELS.get(case.stage, case.stage),
            "approvalTrail": approval_trail,
            "comments": comments,
            "caseId": case.id,
        }

    def escalation_letter_docx_bytes(self, case: EscalationCase) -> bytes:
        letter = self.build_escalation_letter(case)
        if Document is None:
            raise RuntimeError("python-docx is not installed.")
        doc = Document()
        title = doc.add_paragraph()
        run = title.add_run("NIGERIAN MIDSTREAM AND DOWNSTREAM PETROLEUM REGULATORY AUTHORITY (NMDPRA)")
        run.bold = True
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER

        meta = doc.add_table(rows=2, cols=2)
        meta.cell(0, 0).text = f"From: {letter['fromTitle']}"
        meta.cell(0, 1).text = f"Ref: {letter['referenceNumber']}"
        meta.cell(1, 0).text = f"To: {letter['toTitle']}"
        meta.cell(1, 1).text = f"Date: {letter['date']}"

        doc.add_paragraph("")
        subject_p = doc.add_paragraph()
        subject_run = subject_p.add_run(f"Subject: {letter['subject']}")
        subject_run.bold = True
        subject_run.underline = True

        doc.add_paragraph("")
        for para in letter["body"]:
            doc.add_paragraph(para)

        doc.add_paragraph("")
        stage_p = doc.add_paragraph()
        stage_p.add_run(f"Stage of approval: {letter['approvalStage']}").bold = True

        if letter["approvalTrail"]:
            doc.add_paragraph("")
            doc.add_paragraph("Approval / escalation trail:").runs[0].bold = True
            table = doc.add_table(rows=1, cols=4)
            hdr = table.rows[0].cells
            hdr[0].text, hdr[1].text, hdr[2].text, hdr[3].text = "Author", "Stage", "Action", "Timestamp"
            for row in letter["approvalTrail"]:
                cells = table.add_row().cells
                cells[0].text = row["author"]
                cells[1].text = row["stage"]
                cells[2].text = row["action"]
                cells[3].text = row["timestamp"]

        if letter["comments"]:
            doc.add_paragraph("")
            doc.add_paragraph("Comments:").runs[0].bold = True
            for comment in letter["comments"]:
                doc.add_paragraph(comment, style="List Bullet")

        buffer = BytesIO()
        doc.save(buffer)
        return buffer.getvalue()

    def escalation_letter_pdf_bytes(self, case: EscalationCase) -> bytes:
        letter = self.build_escalation_letter(case)
        if pdf_canvas is None:
            raise RuntimeError("reportlab is not installed.")
        buffer = BytesIO()
        c = pdf_canvas.Canvas(buffer, pagesize=A4)
        width, height = A4
        y = height - 25 * mm
        left = 20 * mm

        def line(text: str, size: float = 10.5, bold: bool = False, gap: float = 6.5 * mm) -> None:
            nonlocal y
            c.setFont("Helvetica-Bold" if bold else "Helvetica", size)
            for wrapped in _wrap_text(text, 95):
                if y < 25 * mm:
                    c.showPage()
                    y = height - 25 * mm
                    c.setFont("Helvetica-Bold" if bold else "Helvetica", size)
                c.drawString(left, y, wrapped)
                y -= gap

        c.setFont("Helvetica-Bold", 12)
        c.drawCentredString(width / 2, y, "NIGERIAN MIDSTREAM AND DOWNSTREAM PETROLEUM REGULATORY AUTHORITY")
        y -= 8 * mm
        c.setFont("Helvetica", 9)
        c.drawCentredString(width / 2, y, "(NMDPRA) — Transporter Intelligence")
        y -= 12 * mm

        line(f"From: {letter['fromTitle']}", bold=True)
        line(f"To: {letter['toTitle']}", bold=True)
        line(f"Reference: {letter['referenceNumber']}")
        line(f"Date: {letter['date']}")
        y -= 4 * mm
        line(f"Subject: {letter['subject']}", bold=True)
        y -= 4 * mm
        for para in letter["body"]:
            line(para)
            y -= 2 * mm
        y -= 3 * mm
        line(f"Stage of approval: {letter['approvalStage']}", bold=True)
        if letter["comments"]:
            y -= 3 * mm
            line("Comments:", bold=True)
            for comment in letter["comments"]:
                line(f"- {comment}")
        c.showPage()
        c.save()
        return buffer.getvalue()

    def create_escalation(self, payload: dict[str, Any]) -> dict[str, Any]:
        flag_id = payload.get("flagId")
        flag = self.find_flag(flag_id) if flag_id else None
        if not flag:
            return {"ok": False, "message": "Exception not found."}
        existing = next((case for case in self.escalations.values() if case.flag_id == flag_id), None)
        if existing:
            return {"ok": True, "case": self.escalation_to_view(existing)}
        now = datetime.now(UTC).isoformat()
        case_id = self.next_id("ESC")
        author = clean_text(payload.get("author")) or "Analyst"
        case = EscalationCase(
            case_id,
            flag.id,
            flag.rule,
            flag.severity,
            flag.message,
            flag.supplier_id,
            flag.customer_id,
            flag.affected_date,
            "ANALYST",
            "OPEN",
            self.build_ai_report(flag),
            now,
            now,
            notes=[EscalationNote(author, "ANALYST", "OPENED", "Case opened from AI-flagged exception.", now)],
        )
        self.escalations[case_id] = case
        return {"ok": True, "case": self.escalation_to_view(case)}

    def escalation_action(self, case_id: str, payload: dict[str, Any]) -> dict[str, Any]:
        case = self.escalations.get(case_id)
        if not case:
            return {"ok": False, "message": "Escalation case not found."}
        action = payload.get("action")
        author = clean_text(payload.get("author")) or ESCALATION_STAGE_LABELS.get(case.stage, case.stage)
        text = clean_text(payload.get("note")) or ""
        now = datetime.now(UTC).isoformat()

        if action == "advance":
            idx = ESCALATION_STAGES.index(case.stage) if case.stage in ESCALATION_STAGES else 0
            if idx >= len(ESCALATION_STAGES) - 1:
                return {"ok": False, "message": f"Already at the final stage ({ESCALATION_STAGE_LABELS[case.stage]})."}
            next_stage = ESCALATION_STAGES[idx + 1]
            case.notes.append(EscalationNote(author, case.stage, "ESCALATED", text or f"Escalated to {ESCALATION_STAGE_LABELS[next_stage]}.", now))
            case.stage = next_stage
            case.status = "IN_PROGRESS"
        elif action == "comment":
            if not text:
                return {"ok": False, "message": "Note text is required."}
            case.notes.append(EscalationNote(author, case.stage, "COMMENT", text, now))
        elif action == "resolve":
            case.notes.append(EscalationNote(author, case.stage, "RESOLVED", text or "Marked resolved.", now))
            case.status = "RESOLVED"
        elif action == "close":
            case.notes.append(EscalationNote(author, case.stage, "CLOSED", text or "Case closed.", now))
            case.status = "CLOSED"
        elif action == "reopen":
            case.notes.append(EscalationNote(author, case.stage, "REOPENED", text or "Case reopened.", now))
            case.status = "IN_PROGRESS"
        else:
            return {"ok": False, "message": "Unknown action."}

        case.updated_at = now
        return {"ok": True, "case": self.escalation_to_view(case)}

    def escalation_to_view(self, case: EscalationCase) -> dict[str, Any]:
        result = asdict(case)
        if case.supplier_id and case.supplier_id in self.organisations:
            result["supplier"] = self.organisations[case.supplier_id].name
            result["gasSupplier"] = self.organisations[case.supplier_id].name
            result["shipperId"] = self.shipper_id_for_supplier(case.supplier_id)
            result["shipper"] = self.organisations[result["shipperId"]].name if result["shipperId"] else self.organisations[case.supplier_id].name
        if case.customer_id and case.customer_id in self.organisations:
            result["customer"] = self.organisations[case.customer_id].name
            result["offtaker"] = self.organisations[case.customer_id].name
        result["stageLabel"] = ESCALATION_STAGE_LABELS.get(case.stage, case.stage)
        return result

    def escalations_view(self, filters: dict[str, str]) -> list[dict[str, Any]]:
        stage = filters.get("stage")
        status = filters.get("status")
        cases = list(self.escalations.values())
        if stage:
            cases = [c for c in cases if c.stage == stage]
        if status:
            cases = [c for c in cases if c.status == status]
        cases.sort(key=lambda c: c.updated_at, reverse=True)
        return [self.escalation_to_view(c) for c in cases]

    # ─── Knowledge Base ─────────────────────────────────────────────────────────
    def seed_knowledge_base(self) -> None:
        seed_docs = [
            (
                "Petroleum Industry Act 2021",
                "REGULATION",
                "Principal legislation governing the Nigerian petroleum industry, including "
                "gas transportation network code provisions referenced throughout the platform's "
                "AI Consequence Management engine.",
                ["PIA", "gas transportation", "legislation"],
            ),
            (
                "Nigerian Gas Transportation Network Code (NGTNC)",
                "REGULATION",
                "Network code governing technical, commercial, and operational arrangements for "
                "gas transportation across Nigeria's pipeline network.",
                ["NGTNC", "network code", "balancing"],
            ),
            (
                "NGNR GASCO Data Template — Submission Guide",
                "TEMPLATE",
                "Guidance notes for completing the approved .xlsx/.xlsm workbook template used for "
                "NGIC operator submissions.",
                ["template", "submission", "workbook"],
            ),
        ]
        for title, category, description, tags in seed_docs:
            self.create_knowledge_doc(
                {"title": title, "category": category, "description": description, "tags": ", ".join(tags), "uploadedBy": "NMDPRA Compliance Monitoring"},
                file_bytes=(f"{title}\n\n{description}\n\n(Placeholder reference text — replace by uploading the actual document.)").encode("utf-8"),
                file_name=f"{slug(title)}.txt",
            )

    def create_knowledge_doc(self, payload: dict[str, Any], file_bytes: bytes | None = None, file_name: str | None = None) -> dict[str, Any]:
        title = clean_text(payload.get("title"))
        if not title:
            return {"ok": False, "message": "Document title is required."}
        category = clean_text(payload.get("category")).upper() or "OTHER"
        if category not in KNOWLEDGE_CATEGORIES:
            category = "OTHER"
        tags = [t.strip() for t in clean_text(payload.get("tags")).split(",") if t.strip()]
        doc_id = self.next_id("KB")

        stored_name = None
        mime_type = None
        size_bytes = 0
        if file_bytes is not None and file_name:
            safe_name = re.sub(r"[^A-Za-z0-9_.-]", "_", file_name)
            stored_name = f"{doc_id}-{safe_name}"
            (KNOWLEDGE_BASE_DIR / stored_name).write_bytes(file_bytes)
            mime_type = mimetypes.guess_type(file_name)[0] or "application/octet-stream"
            size_bytes = len(file_bytes)

        doc = KnowledgeDocument(
            doc_id,
            title,
            category,
            clean_text(payload.get("description")),
            tags,
            file_name,
            stored_name,
            mime_type,
            size_bytes,
            clean_text(payload.get("uploadedBy")) or "Unspecified",
            datetime.now(UTC).isoformat(),
        )
        self.knowledge_docs[doc_id] = doc
        return {"ok": True, "doc": self.knowledge_doc_to_view(doc)}

    def delete_knowledge_doc(self, doc_id: str) -> dict[str, Any]:
        doc = self.knowledge_docs.get(doc_id)
        if not doc:
            return {"ok": False, "message": "Document not found."}
        if doc.stored_name:
            target = KNOWLEDGE_BASE_DIR / doc.stored_name
            if target.exists():
                target.unlink()
        del self.knowledge_docs[doc_id]
        return {"ok": True}

    def knowledge_doc_to_view(self, doc: KnowledgeDocument) -> dict[str, Any]:
        result = asdict(doc)
        result["categoryLabel"] = KNOWLEDGE_CATEGORY_LABELS.get(doc.category, doc.category)
        result["hasFile"] = bool(doc.stored_name)
        return result

    def knowledge_docs_view(self, filters: dict[str, str]) -> list[dict[str, Any]]:
        category = filters.get("category")
        query = (filters.get("q") or "").strip().lower()
        docs = list(self.knowledge_docs.values())
        if category:
            docs = [d for d in docs if d.category == category]
        if query:
            docs = [
                d for d in docs
                if query in d.title.lower() or query in d.description.lower() or any(query in t.lower() for t in d.tags)
            ]
        docs.sort(key=lambda d: d.uploaded_at, reverse=True)
        return [self.knowledge_doc_to_view(d) for d in docs]

    # ─── Case Management (manually logged complaints) ──────────────────────────
    def seed_cases(self) -> None:
        seed_rows = [
            (
                "BILLING_ATTRIBUTION",
                "Disputed offtake attribution for March",
                "Egbin Power disputes the attributed offtake volumes recorded against its exit point for the "
                "second half of March, citing a discrepancy with its own metering records.",
                "Adaeze Umeh",
                "Egbin Power",
                "CUSTOMER",
                "a.umeh@egbinpower.example",
                "0803-000-0000",
                None,
                "EGBIN",
                None,
                "HIGH",
                "IN_PROGRESS",
                "Manager, Compliance Monitoring",
            ),
            (
                "GAS_QUALITY",
                "Repeat off-spec condensate carryover",
                "Seplat Energy reports repeat condensate carryover at the Oben entry point affecting downstream "
                "metering accuracy, requesting a joint inspection.",
                "Ifeoma Nwachukwu",
                "Seplat Energy",
                "SUPPLIER",
                "i.nwachukwu@seplat.example",
                "0805-000-0000",
                "SEPLAT",
                None,
                None,
                "MEDIUM",
                "OPEN",
                "Unassigned",
            ),
        ]
        for category, subject, description, name, org, role, email, phone, supplier_id, customer_id, transporter_id, priority, status, assigned_to in seed_rows:
            result = self.create_case({
                "category": category, "subject": subject, "description": description,
                "complainantName": name, "complainantOrg": org, "complainantRole": role,
                "complainantEmail": email, "complainantPhone": phone,
                "relatedSupplierId": supplier_id, "relatedCustomerId": customer_id, "relatedTransporterId": transporter_id,
                "priority": priority, "author": "NMDPRA Compliance Monitoring",
            })
            if result.get("ok") and status != "OPEN":
                case_id = result["case"]["id"]
                if status in ("IN_PROGRESS", "RESOLVED", "CLOSED"):
                    self.case_action(case_id, {"action": "start", "author": assigned_to})
                if status in ("RESOLVED", "CLOSED"):
                    self.case_action(case_id, {"action": "resolve", "author": assigned_to, "note": "Resolved in seed data."})
                if status == "CLOSED":
                    self.case_action(case_id, {"action": "close", "author": assigned_to})
                self.cases[case_id].assigned_to = assigned_to

    def create_case(self, payload: dict[str, Any], file_bytes: bytes | None = None, file_name: str | None = None) -> dict[str, Any]:
        subject = clean_text(payload.get("subject"))
        description = clean_text(payload.get("description"))
        complainant_name = clean_text(payload.get("complainantName"))
        if not subject or not description or not complainant_name:
            return {"ok": False, "message": "Subject, description, and complainant name are required."}
        category = clean_text(payload.get("category")).upper() or "OTHER"
        if category not in COMPLAINT_CATEGORIES:
            category = "OTHER"
        role = clean_text(payload.get("complainantRole")).upper() or "PUBLIC"
        if role not in COMPLAINANT_ROLES:
            role = "PUBLIC"
        priority = clean_text(payload.get("priority")).upper() or "MEDIUM"
        if priority not in COMPLAINT_PRIORITIES:
            priority = "MEDIUM"

        supplier_id = clean_text(payload.get("relatedSupplierId")) or None
        if supplier_id and supplier_id not in self.suppliers:
            supplier_id = None
        customer_id = clean_text(payload.get("relatedCustomerId")) or None
        if customer_id and customer_id not in self.customers:
            customer_id = None
        transporter_id = clean_text(payload.get("relatedTransporterId")) or None
        if transporter_id and transporter_id not in self.transporters:
            transporter_id = None

        case_id = self.next_id("CMP")
        attachment_path = None
        attachment_file_name = None
        if file_bytes is not None and file_name:
            safe_name = re.sub(r"[^A-Za-z0-9_.-]", "_", file_name)
            stored_name = f"{case_id}-{safe_name}"
            (CASE_ATTACHMENTS_DIR / stored_name).write_bytes(file_bytes)
            attachment_path = stored_name
            attachment_file_name = file_name

        now = datetime.now(UTC).isoformat()
        author = clean_text(payload.get("author")) or complainant_name
        case = ComplaintCase(
            case_id,
            category,
            subject,
            description,
            complainant_name,
            clean_text(payload.get("complainantOrg")),
            role,
            clean_text(payload.get("complainantEmail")),
            clean_text(payload.get("complainantPhone")),
            supplier_id,
            customer_id,
            transporter_id,
            priority,
            "OPEN",
            "Unassigned",
            attachment_file_name,
            attachment_path,
            now,
            now,
            notes=[CaseNote(author, "OPENED", "Case logged.", now)],
        )
        self.cases[case_id] = case
        return {"ok": True, "case": self.case_to_view(case)}

    def case_action(self, case_id: str, payload: dict[str, Any]) -> dict[str, Any]:
        case = self.cases.get(case_id)
        if not case:
            return {"ok": False, "message": "Case not found."}
        action = payload.get("action")
        author = clean_text(payload.get("author")) or case.assigned_to or "Analyst"
        text = clean_text(payload.get("note")) or ""
        now = datetime.now(UTC).isoformat()

        if action == "comment":
            if not text:
                return {"ok": False, "message": "Note text is required."}
            case.notes.append(CaseNote(author, "COMMENT", text, now))
        elif action == "assign":
            assignee = clean_text(payload.get("assignedTo"))
            if not assignee:
                return {"ok": False, "message": "Assignee is required."}
            case.assigned_to = assignee
            case.notes.append(CaseNote(author, "ASSIGNED", text or f"Assigned to {assignee}.", now))
            if case.status == "OPEN":
                case.status = "IN_PROGRESS"
        elif action == "start":
            case.status = "IN_PROGRESS"
            case.notes.append(CaseNote(author, "IN_PROGRESS", text or "Investigation started.", now))
        elif action == "resolve":
            case.status = "RESOLVED"
            case.notes.append(CaseNote(author, "RESOLVED", text or "Marked resolved.", now))
        elif action == "close":
            case.status = "CLOSED"
            case.notes.append(CaseNote(author, "CLOSED", text or "Case closed.", now))
        elif action == "reopen":
            case.status = "IN_PROGRESS"
            case.notes.append(CaseNote(author, "REOPENED", text or "Case reopened.", now))
        else:
            return {"ok": False, "message": "Unknown action."}

        case.updated_at = now
        return {"ok": True, "case": self.case_to_view(case)}

    def case_to_view(self, case: ComplaintCase) -> dict[str, Any]:
        result = asdict(case)
        if case.related_supplier_id and case.related_supplier_id in self.organisations:
            result["relatedSupplierName"] = self.organisations[case.related_supplier_id].name
        if case.related_customer_id and case.related_customer_id in self.organisations:
            result["relatedCustomerName"] = self.organisations[case.related_customer_id].name
        if case.related_transporter_id and case.related_transporter_id in self.organisations:
            result["relatedTransporterName"] = self.organisations[case.related_transporter_id].name
        result["categoryLabel"] = COMPLAINT_CATEGORY_LABELS.get(case.category, case.category)
        return result

    def cases_view(self, filters: dict[str, str]) -> list[dict[str, Any]]:
        status = filters.get("status")
        category = filters.get("category")
        priority = filters.get("priority")
        role = filters.get("role")
        cases = list(self.cases.values())
        if status:
            cases = [c for c in cases if c.status == status]
        if category:
            cases = [c for c in cases if c.category == category]
        if priority:
            cases = [c for c in cases if c.priority == priority]
        if role:
            cases = [c for c in cases if c.complainant_role == role]
        cases.sort(key=lambda c: c.updated_at, reverse=True)
        return [self.case_to_view(c) for c in cases]

    # ─── Shipper self-reported nomination (direct entry, cross-checked vs transporter data) ──
    def seed_shipper_nominations(self) -> None:
        seed_rows = [
            ("SHIPPER-SUPPLIER-A", "2026-01-05", 87.5, "Our nomination desk recorded 87.5 MMScf for this gas day; flagging for reconciliation against the transporter's figure."),
            ("SHIPPER-SEPLAT", "2026-01-03", 137.0, "Confirming our nominated volume for this day per our internal scheduling record."),
        ]
        for shipper_id, day, volume, note in seed_rows:
            self.create_shipper_nomination({"shipperId": shipper_id, "date": day, "nominatedVolume": volume, "note": note})

    def create_shipper_nomination(self, payload: dict[str, Any], file_bytes: bytes | None = None, file_name: str | None = None) -> dict[str, Any]:
        shipper_id = clean_text(payload.get("shipperId"))
        if not shipper_id or shipper_id not in self.shippers:
            return {"ok": False, "message": "Unknown shipper."}
        day = clean_text(payload.get("date"))
        if not day:
            return {"ok": False, "message": "Date is required."}
        volume = to_float(payload.get("nominatedVolume"))
        if volume is None:
            return {"ok": False, "message": "Nominated volume is required."}

        nomination_id = self.next_id("SHPNOM")
        stored_name = None
        if file_bytes is not None and file_name:
            safe_name = re.sub(r"[^A-Za-z0-9_.-]", "_", file_name)
            stored_name = f"{nomination_id}-{safe_name}"
            (SHIPPER_NOMINATIONS_DIR / stored_name).write_bytes(file_bytes)

        record = ShipperNomination(
            nomination_id,
            shipper_id,
            day,
            volume,
            clean_text(payload.get("note")),
            datetime.now(UTC).isoformat(),
            file_name,
            stored_name,
        )
        self.shipper_nominations.append(record)
        return {"ok": True, "nomination": self.shipper_nomination_to_view(record)}

    def shipper_nomination_to_view(self, record: ShipperNomination) -> dict[str, Any]:
        result = asdict(record)
        result["hasFile"] = bool(record.stored_name)
        supplier_id = self.supplier_id_for_shipper(record.shipper_id)
        transporter_injection = None
        if supplier_id:
            match = next((r for r in self.supplier_records if r.supplier_id == supplier_id and r.date == record.date), None)
            if match:
                transporter_injection = match.injection_mmscf
        result["transporterInjectionMmscf"] = transporter_injection
        if transporter_injection:
            variance_pct = round(((record.nominated_volume_mmscf - transporter_injection) / transporter_injection) * 100, 2)
            result["variancePct"] = variance_pct
            result["hasDiscrepancy"] = abs(variance_pct) > NOMINATION_VARIANCE_THRESHOLD_PCT
        else:
            result["variancePct"] = None
            result["hasDiscrepancy"] = False
        return result

    def shipper_nominations_view(self, filters: dict[str, str]) -> list[dict[str, Any]]:
        shipper_id = filters.get("shipper")
        records = list(self.shipper_nominations)
        if shipper_id:
            records = [r for r in records if r.shipper_id == shipper_id]
        records.sort(key=lambda r: r.date, reverse=True)
        return [self.shipper_nomination_to_view(r) for r in records]

    def bootstrap(self) -> dict[str, Any]:
        return {
            "organisations": [asdict(org) for org in self.organisations.values()],
            "roles": ROLES,
            "transporters": [self.org_option(org_id) for org_id in self.transporters],
            "shippers": [self.shipper_option(org_id) for org_id in self.shippers],
            "suppliers": [self.org_option(org_id) for org_id in self.suppliers],
            "customers": [self.customer_option(org_id) for org_id in self.customers],
            "sectors": [asdict(sector) for sector in self.sectors.values()],
            "entryPoints": [asdict(point) for point in self.entry_points.values()],
            "exitPoints": [asdict(point) for point in self.exit_points.values()],
            "periods": [asdict(period) for period in self.periods.values()],
            "thresholds": [asdict(t) for t in self.thresholds.values()],
            "exceptionTypes": sorted({flag.rule for flag in self.flags} | {"PRESSURE_OUTSIDE_NEA", "SHRINKAGE_THRESHOLD", "ATTRIBUTION_EXCEEDS_EFFECTIVE_INJECTION", "NOMINATION_VARIANCE", "LINE_PACK_VARIANCE", "MLF_LOSS_THRESHOLD", "OFF_SPEC_GAS"}),
            "monitoringConfig": {
                "nominationVariancePct": NOMINATION_VARIANCE_THRESHOLD_PCT,
                "linePackVariancePct": LINE_PACK_VARIANCE_THRESHOLD_PCT,
                "mlfLossPct": MLF_LOSS_THRESHOLD_PCT,
            },
            "escalationStages": [{"id": s, "label": ESCALATION_STAGE_LABELS[s]} for s in ESCALATION_STAGES],
            "knowledgeCategories": [{"id": c, "label": KNOWLEDGE_CATEGORY_LABELS[c]} for c in KNOWLEDGE_CATEGORIES],
            "complaintCategories": [{"id": c, "label": COMPLAINT_CATEGORY_LABELS[c]} for c in COMPLAINT_CATEGORIES],
            "complaintRoles": COMPLAINANT_ROLES,
            "complaintPriorities": COMPLAINT_PRIORITIES,
            "complaintStatuses": COMPLAINT_STATUSES,
            "scadaPilotActors": SCADA_PILOT_ACTORS,
            "availableYears": sorted({p.start_date[:4] for p in self.periods.values()}, reverse=True),
            "today": to_iso(DEMO_TODAY),
        }

    def org_option(self, org_id: str) -> dict[str, Any]:
        org = self.organisations[org_id]
        return {"id": org.id, "code": org.code, "name": org.name, "roles": org.roles}

    def shipper_option(self, org_id: str) -> dict[str, Any]:
        org = self.organisations[org_id]
        supplier_id = self.shipper_supplier_map.get(org_id)
        supplier = self.organisations.get(supplier_id, Organisation("", "", "", []))
        return {"id": org.id, "code": org.code, "name": org.name, "linkedSupplierId": supplier_id, "linkedSupplier": supplier.name}

    def customer_option(self, org_id: str) -> dict[str, str]:
        org = self.organisations[org_id]
        profile = self.customers[org_id]
        return {"id": org.id, "code": org.code, "name": org.name, "sector": profile.sector_id, "exitPointId": profile.exit_point_id}

    def create_registry_item(self, item_type: str, payload: dict[str, Any]) -> dict[str, Any]:
        code = slug(clean_text(payload.get("code") or payload.get("name") or self.next_id(item_type.upper())))
        name = clean_text(payload.get("name") or code)
        if not name:
            return {"ok": False, "message": "Name is required."}

        if item_type == "transporters":
            self.organisations[code] = Organisation(code, code, name, ["TRANSPORTER"], notes=clean_text(payload.get("notes")))
            self.transporters[code] = TransporterProfile(
                code,
                clean_text(payload.get("networkName")) or f"{name} Network",
                [item.strip() for item in clean_text(payload.get("regions") or "National").split(",") if item.strip()],
                clean_text(payload.get("licenseRef")),
            )
            return {"ok": True, "item": self.org_option(code)}

        if item_type == "suppliers":
            self.organisations[code] = Organisation(code, code, name, ["SUPPLIER"], notes=clean_text(payload.get("notes")))
            entry_point_id = clean_text(payload.get("entryPointId")) or next(iter(self.entry_points))
            entry = self.entry_points.get(entry_point_id) or next(iter(self.entry_points.values()))
            self.suppliers[code] = SupplierProfile(
                code,
                clean_text(payload.get("gasPlant")) or f"{name} Gas Plant",
                to_float(payload.get("plantCapacity")) or 0,
                clean_text(payload.get("meterStreams")) or "1",
                clean_text(payload.get("meterType")) or "Ultrasonic",
                clean_text(payload.get("dgdo")) or "Unassigned",
                to_float(payload.get("dgdoTarget")) or 0,
            )
            self.neas[f"NEA-{code}"] = NetworkEntryAgreement(
                f"NEA-{code}",
                code,
                entry.id,
                clean_text(payload.get("effectiveFrom")) or "2026-01-01",
                clean_text(payload.get("effectiveTo")) or "2026-12-31",
                to_float(payload.get("minPressure")) or entry.min_pressure_barg,
                to_float(payload.get("maxPressure")) or entry.max_pressure_barg,
            )
            self.assignments[f"ASN-NGIC-{code}"] = TransporterSupplierAssignment(f"ASN-NGIC-{code}", "NGIC", code, "2026")
            return {"ok": True, "item": self.org_option(code)}

        if item_type == "customers":
            self.organisations[code] = Organisation(code, code, name, ["CUSTOMER"], notes=clean_text(payload.get("notes")))
            sector_id = clean_text(payload.get("sectorId")) or "GTC"
            exit_point_id = clean_text(payload.get("exitPointId")) or next(iter(self.exit_points))
            self.customers[code] = CustomerProfile(
                code,
                sector_id if sector_id in self.sectors else "GTC",
                exit_point_id if exit_point_id in self.exit_points else next(iter(self.exit_points)),
                to_float(payload.get("contractVolume")) or 0,
                clean_text(payload.get("customerType")) or self.sectors.get(sector_id, self.sectors["GTC"]).name,
            )
            return {"ok": True, "item": self.customer_option(code)}

        if item_type == "entry-points":
            self.entry_points[code] = EntryPoint(
                code,
                code,
                name,
                clean_text(payload.get("location")) or "Unspecified",
                to_float(payload.get("minPressure")) or 0,
                to_float(payload.get("maxPressure")) or 0,
            )
            return {"ok": True, "item": asdict(self.entry_points[code])}

        if item_type == "exit-points":
            self.exit_points[code] = ExitPoint(
                code,
                code,
                name,
                clean_text(payload.get("location")) or "Unspecified",
                to_float(payload.get("minPressure")) or 0,
                to_float(payload.get("maxPressure")) or 0,
                to_float(payload.get("shrinkageThreshold")) or 0.5,
            )
            return {"ok": True, "item": asdict(self.exit_points[code])}

        return {"ok": False, "message": "Unknown registry item type."}

    # ─── Gasco/supplier daily nomination + quality workflow ───────────────────
    def create_gasco_nomination(
        self, payload: dict[str, Any], file_bytes: bytes | None = None, file_name: str | None = None
    ) -> dict[str, Any]:
        supplier_id = clean_text(payload.get("supplierId"))
        if not supplier_id or supplier_id not in self.suppliers:
            return {"ok": False, "message": "Unknown supplier/GASCO."}
        day = clean_text(payload.get("date")) or to_iso(DEMO_TODAY)
        off_spec = bool(payload.get("offSpec"))
        nomination_id = self.next_id("NOM")

        quality_report_path = None
        if file_bytes is not None and file_name:
            safe_name = re.sub(r"[^A-Za-z0-9_.-]", "_", file_name)
            stored_name = f"{nomination_id}-{safe_name}"
            (QUALITY_REPORTS_DIR / stored_name).write_bytes(file_bytes)
            quality_report_path = stored_name

        lab_mlf = to_float(payload.get("labMlf"))
        record = GascoDailyNomination(
            nomination_id,
            supplier_id,
            day,
            to_float(payload.get("previousDayVolume")) or 0.0,
            to_float(payload.get("previousDayInjection")) or 0.0,
            file_name or "not supplied",
            off_spec,
            to_float(payload.get("projectionToday")) or 0.0,
            clean_text(payload.get("disputeNote")),
            datetime.now(UTC).isoformat(),
            quality_report_path,
            lab_mlf,
        )
        self.gasco_nominations.append(record)

        # The quality report's lab-measured MLF feeds straight into the matching day's
        # effective injection figure instead of relying purely on a manually entered MLF.
        if lab_mlf:
            matching_record = next(
                (r for r in self.supplier_records if r.supplier_id == supplier_id and r.date == day),
                None,
            )
            if matching_record is not None:
                matching_record.mlf_mmscf = lab_mlf
                if matching_record.injection_mmscf is not None:
                    matching_record.effective_injection_mmscf = round(matching_record.injection_mmscf * lab_mlf, 4)
        if off_spec:
            submission_id = next(
                (s.id for s in self.submissions.values() if s.supplier_id == supplier_id),
                self.next_id("SUB"),
            )
            flag = self.make_flag(
                submission_id,
                "SUPPLIER_DAY",
                "ERROR",
                "OFF_SPEC_GAS",
                f"Gas quality report for {day} indicates an off-spec delivery from {self.organisations[supplier_id].name}.",
                day,
                supplier_id,
                metric="GAS_QUALITY",
            )
            self.flags.append(flag)
            # Critical potential incidents automatically open a consequence-management
            # escalation case rather than waiting for a manual "Query" click.
            self.create_escalation({"flagId": flag.id, "author": "AI Consequence Management"})
        return {"ok": True, "nomination": asdict(record)}

    def gasco_nominations_view(self, supplier_id: str | None) -> list[dict[str, Any]]:
        rows = [n for n in self.gasco_nominations if not supplier_id or n.supplier_id == supplier_id]
        rows.sort(key=lambda n: n.submitted_at, reverse=True)
        return [asdict(n) for n in rows]

    # ─── Combined injection + volume + quality report (one PDF per supplier/day) ─
    def build_combined_report_pdf(self, supplier_id: str, day: str) -> bytes:
        if pdf_canvas is None:
            raise RuntimeError("reportlab is not installed.")
        supplier_name = self.organisations[supplier_id].name
        supplier_record = next(
            (r for r in self.supplier_records if r.supplier_id == supplier_id and r.date == day), None
        )
        nomination = next(
            (n for n in sorted(self.gasco_nominations, key=lambda n: n.submitted_at, reverse=True) if n.supplier_id == supplier_id and n.date == day),
            None,
        )

        buffer = BytesIO()
        c = pdf_canvas.Canvas(buffer, pagesize=A4)
        width, height = A4
        y = height - 25 * mm
        left = 20 * mm

        def line(text: str, size: float = 10.5, bold: bool = False, gap: float = 6.5 * mm) -> None:
            nonlocal y
            c.setFont("Helvetica-Bold" if bold else "Helvetica", size)
            for wrapped in _wrap_text(text, 95):
                if y < 25 * mm:
                    c.showPage()
                    y = height - 25 * mm
                    c.setFont("Helvetica-Bold" if bold else "Helvetica", size)
                c.drawString(left, y, wrapped)
                y -= gap

        c.setFont("Helvetica-Bold", 12)
        c.drawCentredString(width / 2, y, "NIGERIAN MIDSTREAM AND DOWNSTREAM PETROLEUM REGULATORY AUTHORITY")
        y -= 8 * mm
        c.setFont("Helvetica", 9)
        c.drawCentredString(width / 2, y, "(NMDPRA) — Transporter Intelligence")
        y -= 12 * mm

        line(f"Combined Daily Report — {supplier_name}", size=13, bold=True)
        line(f"Date: {day}")
        y -= 4 * mm

        line("Injection & Transport", size=11.5, bold=True)
        if supplier_record:
            line(f"Entry Pressure: {supplier_record.entry_pressure_barg if supplier_record.entry_pressure_barg is not None else '—'} Barg")
            line(f"Condensate Drop-out: {supplier_record.condensate_ltrs if supplier_record.condensate_ltrs is not None else '—'} Ltrs")
            line(f"Injection: {supplier_record.injection_mmscf if supplier_record.injection_mmscf is not None else '—'} MMScf")
            line(f"MLF: {supplier_record.mlf_mmscf if supplier_record.mlf_mmscf is not None else '—'}")
            line(f"Effective Injection: {supplier_record.effective_injection_mmscf if supplier_record.effective_injection_mmscf is not None else '—'} MMScf")
        else:
            line("No injection record found for this supplier/date.")
        y -= 3 * mm

        line("Nomination & Volume", size=11.5, bold=True)
        if nomination:
            line(f"Previous Day Nomination: {nomination.previous_day_volume_mmscf} MMScf")
            line(f"Previous Day Injection: {nomination.previous_day_injection_mmscf} MMScf")
            line(f"Projection for Today: {nomination.projection_today_mmscf} MMScf")
            if nomination.dispute_note:
                line(f"Dispute Note: {nomination.dispute_note}")
        else:
            line("No nomination submitted for this supplier/date.")
        y -= 3 * mm

        line("Gas Quality", size=11.5, bold=True)
        if nomination:
            line(f"Quality Report: {nomination.quality_report_file_name}")
            line(f"Lab-Measured MLF: {nomination.lab_mlf if nomination.lab_mlf is not None else '—'}")
            line(f"Status: {'OFF-SPEC — fails agreed gas quality specification' if nomination.off_spec else 'Within specification'}", bold=nomination.off_spec)
        else:
            line("No quality report submitted for this supplier/date.")

        c.showPage()
        c.save()
        return buffer.getvalue()

    # ─── Transporter direct daily data entry (no workbook upload) ─────────────
    def create_direct_entry(self, payload: dict[str, Any]) -> dict[str, Any]:
        transporter_id = clean_text(payload.get("transporterId")) or "NGIC"
        supplier_id = clean_text(payload.get("supplierId"))
        if transporter_id not in self.transporters:
            return {"ok": False, "message": "Unknown transporter."}
        if not supplier_id or supplier_id not in self.suppliers:
            return {"ok": False, "message": "Unknown supplier/GASCO."}
        day = clean_text(payload.get("date")) or to_iso(DEMO_TODAY)
        period_id = self.period_for_date(day)
        submission_id = self.next_id("SUB")
        injection = to_float(payload.get("injection"))
        mlf = to_float(payload.get("mlf"))
        effective_injection = round(injection * mlf, 4) if injection is not None and mlf else injection
        record = DailySupplierRecord(
            self.next_id("SREC"),
            submission_id,
            day,
            supplier_id,
            clean_text(payload.get("entryPointId")) or next(iter(self.entry_points)),
            to_float(payload.get("entryPressure")),
            to_float(payload.get("condensate")),
            injection,
            mlf,
            effective_injection,
            0.0,
            0.0,
            0.0,
            0.0,
            None,
            remark="Direct portal data entry (no workbook upload).",
        )
        self.supplier_records.append(record)
        self.submissions[submission_id] = TransporterSubmission(
            submission_id,
            transporter_id,
            supplier_id,
            period_id,
            "DIRECT_ENTRY",
            f"Direct entry — {day}",
            "ACCEPTED",
            datetime.now(UTC).isoformat(),
            {"totalRecords": 1, "acceptedRecords": 1, "warnings": 0, "errors": 0},
        )
        entry = TransporterDirectEntry(
            self.next_id("DE"),
            transporter_id,
            supplier_id,
            day,
            record.entry_pressure_barg,
            record.condensate_ltrs,
            record.injection_mmscf,
            record.mlf_mmscf,
            datetime.now(UTC).isoformat(),
        )
        self.direct_entries.append(entry)
        return {"ok": True, "entry": asdict(entry), "submission": self.submission_to_view(self.submissions[submission_id])}

    def period_for_date(self, day: str) -> str:
        for period in self.periods.values():
            if period.start_date <= day <= period.end_date and period.id not in ("2026", "Q1-2026"):
                return period.id
        return "2026"

    # ─── Daily transporter population status (green/red) ──────────────────────
    def daily_population_status(self, reference_date: str | None = None) -> list[dict[str, Any]]:
        day = reference_date or to_iso(DEMO_TODAY)
        rows = []
        for transporter_id in self.transporters:
            populated = any(
                self.submissions.get(record.submission_id, TransporterSubmission("", "", "", "", "", "", "", "", {})).transporter_id
                == transporter_id
                and record.date == day
                for record in self.supplier_records
            )
            rows.append(
                {
                    "transporterId": transporter_id,
                    "transporter": self.organisations[transporter_id].name,
                    "date": day,
                    "populated": populated,
                    "status": "GREEN" if populated else "RED",
                }
            )
        return rows


STORE = DataStore()


class Handler(BaseHTTPRequestHandler):
    server_version = "TransporterIntelligence/1.0"

    def log_message(self, format: str, *args: Any) -> None:
        return

    def do_GET(self) -> None:
        parsed = urlparse(self.path)
        if parsed.path == "/api/bootstrap":
            self.json_response(STORE.bootstrap())
            return
        if parsed.path == "/api/report":
            params = {key: values[-1] for key, values in parse_qs(parsed.query).items() if values and values[-1]}
            self.json_response(STORE.build_report(params))
            return
        if parsed.path == "/api/escalations":
            params = {key: values[-1] for key, values in parse_qs(parsed.query).items() if values and values[-1]}
            self.json_response({"cases": STORE.escalations_view(params)})
            return
        if parsed.path.startswith("/api/escalations/") and parsed.path.endswith("/letter"):
            case_id = parsed.path.removeprefix("/api/escalations/").removesuffix("/letter")
            case = STORE.escalations.get(case_id)
            if not case:
                self.send_error(404)
                return
            self.json_response(STORE.build_escalation_letter(case))
            return
        if parsed.path.startswith("/api/escalations/") and parsed.path.endswith("/letter.docx"):
            case_id = parsed.path.removeprefix("/api/escalations/").removesuffix("/letter.docx")
            case = STORE.escalations.get(case_id)
            if not case:
                self.send_error(404)
                return
            try:
                data = STORE.escalation_letter_docx_bytes(case)
            except RuntimeError as exc:
                self.json_response({"ok": False, "message": str(exc)}, status=500)
                return
            self.binary_response(data, "application/vnd.openxmlformats-officedocument.wordprocessingml.document", f"{case_id}-escalation-letter.docx")
            return
        if parsed.path.startswith("/api/escalations/") and parsed.path.endswith("/letter.pdf"):
            case_id = parsed.path.removeprefix("/api/escalations/").removesuffix("/letter.pdf")
            case = STORE.escalations.get(case_id)
            if not case:
                self.send_error(404)
                return
            try:
                data = STORE.escalation_letter_pdf_bytes(case)
            except RuntimeError as exc:
                self.json_response({"ok": False, "message": str(exc)}, status=500)
                return
            self.binary_response(data, "application/pdf", f"{case_id}-escalation-letter.pdf")
            return
        if parsed.path == "/api/gasco/nominations":
            params = {key: values[-1] for key, values in parse_qs(parsed.query).items() if values and values[-1]}
            self.json_response({"nominations": STORE.gasco_nominations_view(params.get("supplier"))})
            return
        if parsed.path == "/api/quality-report-file":
            params = {key: values[-1] for key, values in parse_qs(parsed.query).items() if values and values[-1]}
            nomination = next((n for n in STORE.gasco_nominations if n.id == params.get("nominationId")), None)
            if not nomination or not nomination.quality_report_path:
                self.send_error(404)
                return
            target = (QUALITY_REPORTS_DIR / nomination.quality_report_path).resolve()
            if not str(target).startswith(str(QUALITY_REPORTS_DIR.resolve())) or not target.exists():
                self.send_error(404)
                return
            self.binary_response(target.read_bytes(), "application/pdf", nomination.quality_report_file_name or "quality-report.pdf")
            return
        if parsed.path == "/api/combined-report":
            params = {key: values[-1] for key, values in parse_qs(parsed.query).items() if values and values[-1]}
            supplier_id, day = params.get("supplierId"), params.get("date")
            if not supplier_id or supplier_id not in STORE.suppliers or not day:
                self.send_error(404)
                return
            try:
                data = STORE.build_combined_report_pdf(supplier_id, day)
            except RuntimeError as exc:
                self.json_response({"ok": False, "message": str(exc)}, status=500)
                return
            self.binary_response(data, "application/pdf", f"{supplier_id}-{day}-combined-report.pdf")
            return
        if parsed.path == "/api/shipper/nominations":
            params = {key: values[-1] for key, values in parse_qs(parsed.query).items() if values and values[-1]}
            self.json_response({"nominations": STORE.shipper_nominations_view(params)})
            return
        if parsed.path.startswith("/api/shipper/nominations/") and parsed.path.endswith("/file"):
            nom_id = parsed.path.removeprefix("/api/shipper/nominations/").removesuffix("/file")
            record = next((r for r in STORE.shipper_nominations if r.id == nom_id), None)
            if not record or not record.stored_name:
                self.send_error(404)
                return
            target = (SHIPPER_NOMINATIONS_DIR / record.stored_name).resolve()
            if not str(target).startswith(str(SHIPPER_NOMINATIONS_DIR.resolve())) or not target.exists():
                self.send_error(404)
                return
            mime_type = mimetypes.guess_type(record.file_name or "")[0] or "application/octet-stream"
            self.binary_response(target.read_bytes(), mime_type, record.file_name or "nomination-file")
            return
        if parsed.path == "/api/knowledge-base":
            params = {key: values[-1] for key, values in parse_qs(parsed.query).items() if values and values[-1]}
            self.json_response({"docs": STORE.knowledge_docs_view(params)})
            return
        if parsed.path.startswith("/api/knowledge-base/") and parsed.path.endswith("/file"):
            doc_id = parsed.path.removeprefix("/api/knowledge-base/").removesuffix("/file")
            doc = STORE.knowledge_docs.get(doc_id)
            if not doc or not doc.stored_name:
                self.send_error(404)
                return
            target = (KNOWLEDGE_BASE_DIR / doc.stored_name).resolve()
            if not str(target).startswith(str(KNOWLEDGE_BASE_DIR.resolve())) or not target.exists():
                self.send_error(404)
                return
            self.binary_response(target.read_bytes(), doc.mime_type or "application/octet-stream", doc.file_name or doc.stored_name)
            return
        if parsed.path == "/api/cases":
            params = {key: values[-1] for key, values in parse_qs(parsed.query).items() if values and values[-1]}
            self.json_response({"cases": STORE.cases_view(params)})
            return
        if parsed.path.startswith("/api/cases/") and parsed.path.endswith("/attachment"):
            case_id = parsed.path.removeprefix("/api/cases/").removesuffix("/attachment")
            case = STORE.cases.get(case_id)
            if not case or not case.attachment_path:
                self.send_error(404)
                return
            target = (CASE_ATTACHMENTS_DIR / case.attachment_path).resolve()
            if not str(target).startswith(str(CASE_ATTACHMENTS_DIR.resolve())) or not target.exists():
                self.send_error(404)
                return
            mime_type = mimetypes.guess_type(case.attachment_file_name or "")[0] or "application/octet-stream"
            self.binary_response(target.read_bytes(), mime_type, case.attachment_file_name or "attachment")
            return
        if parsed.path == "/api/population-status":
            params = {key: values[-1] for key, values in parse_qs(parsed.query).items() if values and values[-1]}
            self.json_response({"rows": STORE.daily_population_status(params.get("date"))})
            return
        path = "index.html" if parsed.path in {"/", ""} else parsed.path.lstrip("/")
        self.serve_static(path)

    def do_POST(self) -> None:
        parsed = urlparse(self.path)
        if parsed.path == "/api/escalations":
            length = int(self.headers.get("Content-Length", "0") or 0)
            try:
                payload = json.loads(self.rfile.read(length).decode("utf-8") or "{}")
            except json.JSONDecodeError:
                self.json_response({"ok": False, "message": "Invalid JSON."}, status=400)
                return
            result = STORE.create_escalation(payload)
            self.json_response(result, status=200 if result.get("ok") else 422)
            return

        if parsed.path.startswith("/api/escalations/") and parsed.path.endswith("/action"):
            case_id = parsed.path.removeprefix("/api/escalations/").removesuffix("/action")
            length = int(self.headers.get("Content-Length", "0") or 0)
            try:
                payload = json.loads(self.rfile.read(length).decode("utf-8") or "{}")
            except json.JSONDecodeError:
                self.json_response({"ok": False, "message": "Invalid JSON."}, status=400)
                return
            result = STORE.escalation_action(case_id, payload)
            self.json_response(result, status=200 if result.get("ok") else 422)
            return

        if parsed.path.startswith("/api/thresholds/"):
            rule_id = parsed.path.removeprefix("/api/thresholds/")
            length = int(self.headers.get("Content-Length", "0") or 0)
            try:
                payload = json.loads(self.rfile.read(length).decode("utf-8") or "{}")
            except json.JSONDecodeError:
                self.json_response({"ok": False, "message": "Invalid JSON."}, status=400)
                return
            self.json_response(STORE.update_threshold(rule_id, payload))
            return

        if parsed.path == "/api/gasco/nominations":
            form = cgi.FieldStorage(
                fp=self.rfile,
                headers=self.headers,
                environ={"REQUEST_METHOD": "POST", "CONTENT_TYPE": self.headers.get("Content-Type", "")},
            )
            payload = {
                "supplierId": form.getfirst("supplierId", ""),
                "date": form.getfirst("date", ""),
                "previousDayVolume": form.getfirst("previousDayVolume", ""),
                "previousDayInjection": form.getfirst("previousDayInjection", ""),
                "labMlf": form.getfirst("labMlf", ""),
                "offSpec": form.getfirst("offSpec", ""),
                "projectionToday": form.getfirst("projectionToday", ""),
                "disputeNote": form.getfirst("disputeNote", ""),
            }
            file_item = form["qualityReportFile"] if "qualityReportFile" in form else None
            file_bytes = None
            file_name = None
            if file_item is not None and getattr(file_item, "filename", ""):
                file_bytes = file_item.file.read()
                file_name = file_item.filename
            result = STORE.create_gasco_nomination(payload, file_bytes, file_name)
            self.json_response(result, status=200 if result.get("ok") else 422)
            return

        if parsed.path == "/api/shipper/nominations":
            form = cgi.FieldStorage(
                fp=self.rfile,
                headers=self.headers,
                environ={"REQUEST_METHOD": "POST", "CONTENT_TYPE": self.headers.get("Content-Type", "")},
            )
            payload = {
                "shipperId": form.getfirst("shipperId", ""),
                "date": form.getfirst("date", ""),
                "nominatedVolume": form.getfirst("nominatedVolume", ""),
                "note": form.getfirst("note", ""),
            }
            file_item = form["file"] if "file" in form else None
            file_bytes = None
            file_name = None
            if file_item is not None and getattr(file_item, "filename", ""):
                file_bytes = file_item.file.read()
                file_name = file_item.filename
            result = STORE.create_shipper_nomination(payload, file_bytes, file_name)
            self.json_response(result, status=200 if result.get("ok") else 422)
            return

        if parsed.path == "/api/knowledge-base":
            form = cgi.FieldStorage(
                fp=self.rfile,
                headers=self.headers,
                environ={"REQUEST_METHOD": "POST", "CONTENT_TYPE": self.headers.get("Content-Type", "")},
            )
            payload = {
                "title": form.getfirst("title", ""),
                "category": form.getfirst("category", ""),
                "description": form.getfirst("description", ""),
                "tags": form.getfirst("tags", ""),
                "uploadedBy": form.getfirst("uploadedBy", ""),
            }
            file_item = form["file"] if "file" in form else None
            file_bytes = None
            file_name = None
            if file_item is not None and getattr(file_item, "filename", ""):
                file_bytes = file_item.file.read()
                file_name = file_item.filename
            result = STORE.create_knowledge_doc(payload, file_bytes, file_name)
            self.json_response(result, status=200 if result.get("ok") else 422)
            return

        if parsed.path.startswith("/api/knowledge-base/") and parsed.path.endswith("/delete"):
            doc_id = parsed.path.removeprefix("/api/knowledge-base/").removesuffix("/delete")
            result = STORE.delete_knowledge_doc(doc_id)
            self.json_response(result, status=200 if result.get("ok") else 404)
            return

        if parsed.path == "/api/cases":
            form = cgi.FieldStorage(
                fp=self.rfile,
                headers=self.headers,
                environ={"REQUEST_METHOD": "POST", "CONTENT_TYPE": self.headers.get("Content-Type", "")},
            )
            payload = {
                "category": form.getfirst("category", ""),
                "subject": form.getfirst("subject", ""),
                "description": form.getfirst("description", ""),
                "complainantName": form.getfirst("complainantName", ""),
                "complainantOrg": form.getfirst("complainantOrg", ""),
                "complainantRole": form.getfirst("complainantRole", ""),
                "complainantEmail": form.getfirst("complainantEmail", ""),
                "complainantPhone": form.getfirst("complainantPhone", ""),
                "relatedSupplierId": form.getfirst("relatedSupplierId", ""),
                "relatedCustomerId": form.getfirst("relatedCustomerId", ""),
                "relatedTransporterId": form.getfirst("relatedTransporterId", ""),
                "priority": form.getfirst("priority", ""),
                "author": form.getfirst("author", ""),
            }
            file_item = form["attachment"] if "attachment" in form else None
            file_bytes = None
            file_name = None
            if file_item is not None and getattr(file_item, "filename", ""):
                file_bytes = file_item.file.read()
                file_name = file_item.filename
            result = STORE.create_case(payload, file_bytes, file_name)
            self.json_response(result, status=200 if result.get("ok") else 422)
            return

        if parsed.path.startswith("/api/cases/") and parsed.path.endswith("/action"):
            case_id = parsed.path.removeprefix("/api/cases/").removesuffix("/action")
            length = int(self.headers.get("Content-Length", "0") or 0)
            try:
                payload = json.loads(self.rfile.read(length).decode("utf-8") or "{}")
            except json.JSONDecodeError:
                self.json_response({"ok": False, "message": "Invalid JSON."}, status=400)
                return
            result = STORE.case_action(case_id, payload)
            self.json_response(result, status=200 if result.get("ok") else 422)
            return

        if parsed.path == "/api/daily-entry":
            length = int(self.headers.get("Content-Length", "0") or 0)
            try:
                payload = json.loads(self.rfile.read(length).decode("utf-8") or "{}")
            except json.JSONDecodeError:
                self.json_response({"ok": False, "message": "Invalid JSON."}, status=400)
                return
            result = STORE.create_direct_entry(payload)
            self.json_response(result, status=200 if result.get("ok") else 422)
            return

        if parsed.path.startswith("/api/registry/"):
            item_type = parsed.path.removeprefix("/api/registry/")
            length = int(self.headers.get("Content-Length", "0") or 0)
            try:
                payload = json.loads(self.rfile.read(length).decode("utf-8") or "{}")
            except json.JSONDecodeError:
                self.json_response({"ok": False, "message": "Invalid JSON payload."}, status=400)
                return
            result = STORE.create_registry_item(item_type, payload)
            self.json_response(result, status=200 if result.get("ok") else 422)
            return

        if parsed.path != "/api/uploads":
            self.send_error(404)
            return
        form = cgi.FieldStorage(
            fp=self.rfile,
            headers=self.headers,
            environ={"REQUEST_METHOD": "POST", "CONTENT_TYPE": self.headers.get("Content-Type", "")},
        )
        file_item = form["file"] if "file" in form else None
        if file_item is None or not getattr(file_item, "file", None):
            self.json_response({"ok": False, "errors": [{"message": "No workbook file was provided."}], "warnings": []}, status=400)
            return
        transporter_id = form.getfirst("transporterId", "") or "NGIC"
        supplier_id = form.getfirst("supplierId", "") or ""
        period_id = form.getfirst("periodId", "") or ""
        file_name = getattr(file_item, "filename", "upload.xlsx")
        result = STORE.parse_workbook(file_item.file.read(), file_name, transporter_id, supplier_id, period_id)
        self.json_response(result, status=200 if result["ok"] else 422)

    def serve_static(self, path: str) -> None:
        target = (STATIC_DIR / path).resolve()
        if not str(target).startswith(str(STATIC_DIR.resolve())) or not target.exists() or target.is_dir():
            self.send_error(404)
            return
        content = target.read_bytes()
        mime_type = mimetypes.guess_type(str(target))[0] or "application/octet-stream"
        self.send_response(200)
        self.send_header("Content-Type", mime_type)
        self.send_header("Content-Length", str(len(content)))
        self.end_headers()
        self.wfile.write(content)

    def binary_response(self, data: bytes, mime_type: str, file_name: str) -> None:
        self.send_response(200)
        self.send_header("Content-Type", mime_type)
        self.send_header("Content-Length", str(len(data)))
        self.send_header("Content-Disposition", f'attachment; filename="{file_name}"')
        self.end_headers()
        self.wfile.write(data)

    def json_response(self, payload: Any, status: int = 200) -> None:
        encoded = json.dumps(payload, default=str).encode("utf-8")
        self.send_response(status)
        self.send_header("Content-Type", "application/json")
        self.send_header("Content-Length", str(len(encoded)))
        self.end_headers()
        self.wfile.write(encoded)


def main() -> None:
    host = os.environ.get("HOST", "0.0.0.0")
    port = int(os.environ.get("PORT", "8765"))
    server = ThreadingHTTPServer((host, port), Handler)
    print(f"Transporter Intelligence running at http://127.0.0.1:{port} (bound on {host}:{port})")
    server.serve_forever()


if __name__ == "__main__":
    main()
